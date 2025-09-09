import os
import io
import base64
import sqlite3
import subprocess
import shutil
import sys
import importlib
import re
import uuid
import zipfile
from datetime import datetime, timedelta
from typing import Dict, Optional, List, Tuple
import streamlit as st
import streamlit.components.v1 as components
import pandas as pd
from docx import Document
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.text.run import Run
from docx.oxml.shared import OxmlElement, qn
from docx.document import Document as DocxDocument
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.shared import Inches
from docx2pdf import convert as docx2pdf_convert

# ---------------------------
# Constants and configuration
# ---------------------------
APP_TITLE = "Mierae Invoice/Quotation Manager"
DB_PATH = os.path.join(os.getcwd(), "invoices.db")
TEMPLATE_PATH = os.path.join(os.getcwd(), "Mierae Quotation Template new.docx")
TEMPLATE_PATH_55 = os.path.join(os.getcwd(), "mierae quotation 5.4.docx")
OUTPUT_DIR = os.path.join(os.getcwd(), "output")
DOCX_DIR = os.path.join(OUTPUT_DIR, "docx")
PDF_DIR = os.path.join(OUTPUT_DIR, "pdf")
# New: Agreement/Feasibility storage
AGREEMENT_DIR = os.path.join(OUTPUT_DIR, "agreements")
FEASIBILITY_DIR = os.path.join(AGREEMENT_DIR, "feasibility")
AGREEMENT_PDF_DIR = os.path.join(AGREEMENT_DIR, "pdf")

# Custom exceptions
class DuplicateAgreementError(Exception):
    pass

PRODUCT_OPTIONS = [
    "3.3 kW Residential Rooftop Solar System",
    "5.5 kW Residential Rooftop Solar System",
]

QUOTATION_PREFIX = "MIERAE/25-26/"
QUOTATION_START_NUMBER = 793  # corresponds to 0001

# ---------------------------
# Utility: database
# ---------------------------

def get_conn() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH)
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS invoices (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            quotation_no TEXT UNIQUE NOT NULL,
            product TEXT NOT NULL,
            customer_name TEXT NOT NULL,
            mobile TEXT,
            location TEXT,
            city TEXT,
            state TEXT,
            pincode TEXT,
            staff_name TEXT,
            date_of_quotation TEXT,
            validity_date TEXT,
            docx_path TEXT,
            pdf_path TEXT,
            created_at TEXT,
            updated_at TEXT
        )
        """
    )
    # New: agreements table to store feasibility uploads and generated agreements
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS agreements (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            agreement_no TEXT UNIQUE NOT NULL,
            name TEXT,
            number TEXT,
            address TEXT,
            date TEXT,
            feasibility_pdf_path TEXT,
            agreement_pdf_path TEXT,
            created_at TEXT
        )
        """
    )
    # Log of feasibility uploads (to avoid relying on filesystem for counts)
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS feasibility_events (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            uploaded_at TEXT NOT NULL
        )
        """
    )
    return conn

# ---------------------------
# Helpers for quotation number
# ---------------------------

def next_quotation_no(conn: sqlite3.Connection) -> str:
    cur = conn.cursor()
    cur.execute(
        "SELECT quotation_no FROM invoices WHERE quotation_no LIKE ? ORDER BY id DESC LIMIT 1",
        (f"{QUOTATION_PREFIX}%",),
    )
    row = cur.fetchone()
    if row and isinstance(row[0], str):
        try:
            suffix = int(row[0].split("/")[-1])
            nxt = suffix + 1
        except Exception:
            nxt = QUOTATION_START_NUMBER
    else:
        nxt = QUOTATION_START_NUMBER
    return f"{QUOTATION_PREFIX}{nxt:04d}"

# ---------------------------
# File system helpers
# ---------------------------

def ensure_dirs():
    os.makedirs(DOCX_DIR, exist_ok=True)
    os.makedirs(PDF_DIR, exist_ok=True)
    # New: ensure agreement-related directories
    os.makedirs(AGREEMENT_DIR, exist_ok=True)
    os.makedirs(FEASIBILITY_DIR, exist_ok=True)
    os.makedirs(AGREEMENT_PDF_DIR, exist_ok=True)

# ---------------------------
# Helper: upload PDF to transfer.sh
# ---------------------------

def _upload_to_transfersh(file_path: str) -> Optional[str]:
    """Uploads a file to transfer.sh via HTTP PUT and returns the public URL, or None on failure.
    Uses only stdlib to avoid new dependencies.
    """
    try:
        import urllib.request as _ur
        filename = os.path.basename(file_path)
        url = f"https://transfer.sh/{filename}"
        req = _ur.Request(url, method="PUT")
        with open(file_path, "rb") as f:
            data = f.read()
        req.add_header("Content-Type", "application/octet-stream")
        req.add_header("Content-Length", str(len(data)))
        with _ur.urlopen(req, data=data, timeout=60) as resp:
            body = resp.read().decode().strip()
            # transfer.sh usually responds with the final URL in the body
            if body.startswith("http://") or body.startswith("https://"):
                return body
            # Fallback to request URL if a 200 with no body URL
            return url
    except Exception:
        return None


# ---------------------------
# Helpers for sharing/public links and mobile share UI
# ---------------------------

def _get_secret(key: str, default: Optional[str] = None) -> Optional[str]:
    """Safe accessor for Streamlit secrets. Returns default if missing."""
    try:
        import streamlit as _st
        if hasattr(_st, "secrets") and key in _st.secrets:
            val = _st.secrets.get(key)
            return str(val) if val is not None else default
    except Exception:
        pass


def _render_delete_button(record_id: int, label: str = "Delete") -> None:
    """Render a red pill 'Delete' button via HTML/JS that triggers deletion using query params.
    We redirect to the same page with ?delete_id=ID, which is handled in Python to perform deletion.
    """
    try:
        import streamlit.components.v1 as _components
        btn_html = f"""
        <div style=\"display:flex; justify-content:center; width:100%\">
          <button
            onclick=\"(function(){{ const u=new URL(window.location.href); u.searchParams.set('delete_id','{record_id}'); u.searchParams.set('ts', Date.now()); window.location.href=u.toString(); }})()\"
            style=\"
              display:inline-flex; align-items:center; gap:8px;
              padding: 0.35rem 0.6rem; border-radius:999px;
              font-size:13px; min-width:36px; height:36px;
              background:#ef4444; color:#ffffff; border:none; cursor:pointer;
              box-shadow: 0 1px 2px rgba(0,0,0,0.05);
            \">
            <span>🗑️</span>
            <span>{label}</span>
          </button>
        </div>
        """
        _components.html(btn_html, height=60)
    except Exception:
        pass


def _handle_delete_via_query() -> None:
    """Check URL query params for delete_id and delete the record if present."""
    try:
        import streamlit as _st
        del_id = _st.query_params.get("delete_id")
        if del_id:
            try:
                rid = int(del_id if isinstance(del_id, str) else del_id[0])
                delete_invoice(rid)
                _st.success("Deleted.")
            except Exception as e:
                _st.error(f"Failed to delete: {e}")
            finally:
                # Clear the param and rerun to refresh the list
                try:
                    _st.query_params.clear()
                except Exception:
                    pass
                _st.rerun()
    except Exception:
        pass



def _upload_to_fileio(file_path: str) -> Optional[str]:
    """Upload a file to file.io and return a public URL. Uses stdlib only.
    Note: file.io links may expire by default. This is a best-effort fallback.
    """
    try:
        import os as _os
        import json as _json
        import uuid as _uuid
        import urllib.request as _ur

        boundary = f"----WebKitFormBoundary{_uuid.uuid4().hex}"
        filename = _os.path.basename(file_path)
        with open(file_path, "rb") as f:
            file_bytes = f.read()

        # Build multipart/form-data body
        parts = []
        parts.append(f"--{boundary}\r\n".encode())
        parts.append(
            (
                f"Content-Disposition: form-data; name=\"file\"; filename=\"{filename}\"\r\n"
                f"Content-Type: application/pdf\r\n\r\n"
            ).encode()
        )
        parts.append(file_bytes)
        parts.append("\r\n".encode())
        # Optional: set maxDownloads=1 or expiry; we omit to keep default
        parts.append(f"--{boundary}--\r\n".encode())
        body = b"".join(parts)

        req = _ur.Request("https://file.io")
        req.add_header("Content-Type", f"multipart/form-data; boundary={boundary}")
        req.add_header("Content-Length", str(len(body)))
        with _ur.urlopen(req, data=body, timeout=60) as resp:
            raw = resp.read().decode("utf-8", errors="ignore")
            try:
                data = _json.loads(raw)
                url = data.get("link") or data.get("url") or data.get("success")
                if isinstance(url, str) and (url.startswith("http://") or url.startswith("https://")):
                    return url
            except Exception:
                # Some responses are plain text URL
                if raw.startswith("http://") or raw.startswith("https://"):
                    return raw.strip()
        return None
    except Exception:
        return None


def _get_public_link(file_path: str) -> Optional[str]:
    """Return a cached public URL for the file, uploading if needed.
    Tries transfer.sh first, falls back to file.io. Caches by absolute path.
    """
    try:
        import os as _os
        import streamlit as _st
        key = f"public_url::{_os.path.abspath(file_path)}"
        cached = _st.session_state.get(key)
        if cached:
            return cached
        # Try transfer.sh
        url = _upload_to_transfersh(file_path)
        if not url:
            url = _upload_to_fileio(file_path)
        if url:
            _st.session_state[key] = url
        return url
    except Exception:
        return None


def _render_mobile_share_button(pdf_path: str, filename: Optional[str] = None) -> None:
    """Render a mobile-friendly Share button that shares the actual PDF file via Web Share API.
    Falls back to a normal download link if file sharing is not supported.
    """
    try:
        import os as _os
        import base64 as _b64
        import streamlit as _st
        import streamlit.components.v1 as _components

        if not (pdf_path and _os.path.exists(pdf_path)):
            _st.warning("PDF not found for sharing.")
            return
        name = filename or _os.path.basename(pdf_path) or "invoice.pdf"
        with open(pdf_path, "rb") as f:
            b64 = _b64.b64encode(f.read()).decode("utf-8")

        html = f"""
        <div style="display:flex; justify-content:center; width:100%">
          <button id="sharePdfBtn" aria-label="Share PDF"
                  style="
                    display:inline-flex; align-items:center; gap:8px;
                    padding: 0.35rem 0.6rem; border-radius:999px;
                    font-size:13px; min-width:36px; height:36px;
                    background:#10b981; color:#ffffff; border:none; cursor:pointer;
                    box-shadow: 0 1px 2px rgba(0,0,0,0.05);
                  ">
            <span>📤</span>
            <span>Share PDF</span>
          </button>
          <a id="dlLink" download="{name}" href="data:application/pdf;base64,{b64}" style="display:none">Download</a>
        </div>
        <script>
        (function() {{
          const btn = document.getElementById('sharePdfBtn');
          const dl = document.getElementById('dlLink');
          const b64 = "{b64}";
          const fname = "{name}";
          function b64ToBytes(b64) {{
            const bin = atob(b64);
            const len = bin.length;
            const bytes = new Uint8Array(len);
            for (let i = 0; i < len; i++) bytes[i] = bin.charCodeAt(i);
            return bytes;
          }}
          btn.addEventListener('click', async () => {{
            try {{
              const bytes = b64ToBytes(b64);
              const blob = new Blob([bytes], {{ type: 'application/pdf' }});
              const file = new File([blob], fname, {{ type: 'application/pdf' }});
              if (navigator.canShare && navigator.canShare({{ files: [file] }})) {{
                await navigator.share({{
                  files: [file],
                  title: fname,
                  text: 'Invoice PDF'
                }});
              }} else {{
                // Fallback: trigger download so user can attach from gallery/files
                dl.click();
              }}
            }} catch (e) {{
              console.error(e);
              dl.click();
            }}
          }});
        }})();
        </script>
        """
        _components.html(html, height=60)
    except Exception:
        # Quietly ignore; UI fallback handled elsewhere
        pass


def safe_filename(name: str) -> str:
    """Return a filesystem-safe filename fragment (no path separators or illegal chars)."""
    illegal = ['\\\\', '/', ':', '*', '?', '"', '<', '>', '|']
    safe = name
    for ch in illegal:
        safe = safe.replace(ch, '-')
    # Collapse spaces and trim
    safe = " ".join(safe.split())
    return safe

# ---------------------------
# Helper: WhatsApp Cloud API send (optional)
# ---------------------------

def _wa_send_document_via_cloud_api(
    phone_number: str,
    doc_link: str,
    filename: str,
    caption: str,
    token: str,
    phone_number_id: str,
) -> Tuple[bool, str]:
    """Send a document message using WhatsApp Cloud API. Returns (ok, msg)."""
    try:
        import json as _json
        import urllib.request as _ur
        import urllib.error as _ue

        url = f"https://graph.facebook.com/v20.0/{phone_number_id}/messages"
        payload = {
            "messaging_product": "whatsapp",
            "to": phone_number,
            "type": "document",
            "document": {
                "link": doc_link,
                "filename": filename,
                "caption": caption or "",
            },
        }
        data = _json.dumps(payload).encode("utf-8")
        req = _ur.Request(url, data=data, method="POST")
        req.add_header("Authorization", f"Bearer {token}")
        req.add_header("Content-Type", "application/json")
        try:
            with _ur.urlopen(req, timeout=60) as resp:
                _ = resp.read()
                return True, "Sent via WhatsApp API."
        except _ue.HTTPError as e:
            try:
                err_body = e.read().decode()
            except Exception:
                err_body = str(e)
            return False, f"API error: {err_body}"
    except Exception as ex:
        return False, f"Failed: {ex}"

# ---------------------------
# DOCX processing - replace only yellow highlighted runs
# ---------------------------

def iter_paragraphs_and_cells(doc: DocxDocument) -> List[Paragraph]:
    items: List[Paragraph] = []
    # Paragraphs at document level
    items.extend(doc.paragraphs)
    # Paragraphs inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                items.extend(cell.paragraphs)
                # Also nested tables inside a cell
                for tbl in cell.tables:
                    for r in tbl.rows:
                        for c in r.cells:
                            items.extend(c.paragraphs)
    return items


def get_yellow_runs(doc: DocxDocument) -> List[Run]:
    yellow_runs: List[Run] = []
    for p in iter_paragraphs_and_cells(doc):
        for run in p.runs:
            try:
                if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                    yellow_runs.append(run)
            except Exception:
                # Some runs might not have highlight attribute accessible
                pass
    return yellow_runs


def replace_yellow_fields(doc: DocxDocument, values_in_order: List[str]) -> None:
    """Backward-compatible: keep for potential future use."""
    runs = get_yellow_runs(doc)
    for i, run in enumerate(runs):
        if i < len(values_in_order):
            run.text = str(values_in_order[i])


def replace_by_labels(doc: DocxDocument, data: Dict[str, str]) -> None:
    """Replace highlighted values based on paragraph labels to avoid misaligned fields.
    Labels handled:
    - Customer Name:
    - Location:
    - City:
    - State:
    - Phone:
    - Product & Service:
    - Quotation No:
    - Date of Quotation:
    - Validity of Quotation:
    """
    def fmt_date(val: str) -> str:
        try:
            # Accept YYYY-MM-DD and return DD/MM/YYYY
            dt = datetime.strptime(val, "%Y-%m-%d")
            return dt.strftime("%d/%m/%Y")
        except Exception:
            return val or ""

    targets = {
        "customer name": data.get("customer_name", ""),
        "location": data.get("location", "").replace("\n", " ").strip(),
        "city": data.get("city", ""),
        "state": data.get("state", ""),
        # Explicit pincode labels (handle both single word and spaced variant)
        "pincode": data.get("pincode", ""),
        "pin code": data.get("pincode", ""),
        # Support multiple possible labels used in templates for phone/mobile/customer no
        "phone": data.get("mobile", ""),
        "customer no": data.get("mobile", ""),
        "mobile no": data.get("mobile", ""),
        "mobile number": data.get("mobile", ""),
        "product & service": data.get("product", ""),
        "date of quotation": fmt_date(data.get("date_of_quotation", "")),
        "validity of quotation": fmt_date(data.get("validity_date", "")),
        "quotation no": str(data.get("quotation_no", "")),
    }

    # Labels for which we should remove the title text (customer info block only)
    STRIP_LABELS = set([
        "customer name", "location", "city", "state",
        "pincode", "pin code",
        "phone", "customer no", "mobile no", "mobile number",
    ])

    def replace_in_paragraph(p: Paragraph, label: str, value: str, all_labels: List[str]):
        text = p.text
        text_lower = text.lower()
        # detect label with ':' or '-'
        candidates = [f"{label}:", f"{label}-"]
        label_start = -1
        label_end = -1
        for cand in candidates:
            idx = text_lower.find(cand)
            if idx != -1:
                label_start = idx
                label_end = idx + len(cand)
                break
        if label_start == -1:
            return  # label not in this paragraph

        # find the next other label occurrence to bound our clearing range
        next_idx = len(text)
        for other in all_labels:
            if other == label:
                continue
            for sep in (":", "-"):
                i = text_lower.find(f"{other}{sep}", label_end)
                if i != -1:
                    next_idx = min(next_idx, i)
        # iterate runs and find first yellow run whose run range begins after label_end and before next_idx
        pos = 0
        replaced = False
        for r in p.runs:
            rt = r.text
            begin = pos
            end = pos + len(rt)
            pos = end
            if end <= label_end:
                continue
            if begin >= next_idx:
                break
            try:
                if r.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                    if not replaced:
                        r.text = "" if value is None else str(value)
                        replaced = True
                    else:
                        # clear leftover highlighted placeholders in this label's region
                        r.text = ""
            except Exception:
                pass

        # Second pass: remove the label text portion itself ONLY for customer info labels
        if label in STRIP_LABELS:
            pos = 0
            for r in p.runs:
                rt = r.text
                begin = pos
                end = pos + len(rt)
                pos = end
                # full overlap with label => clear
                if end <= label_end and end > label_start:
                    r.text = ""
                # partial overlap => trim the label part
                elif begin < label_end < end:
                    keep_from = label_end - begin
                    try:
                        r.text = rt[keep_from:]
                    except Exception:
                        pass

    search_paras = iter_paragraphs_and_cells(doc)
    label_list = list(targets.keys())
    for p in search_paras:
        # Combined State and Pincode in same paragraph special-case still supported
        tl = p.text.lower()
        if ("state:" in tl or "state-" in tl) and ("pincode" in tl or "pin code" in tl):
            values = [data.get("state", ""), data.get("pincode", "")]
            idx = 0
            for r in p.runs:
                try:
                    if r.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                        if idx < len(values):
                            r.text = str(values[idx])
                            idx += 1
                except Exception:
                    pass
            # If we only managed to set State (one yellow run) but not Pincode, try to place pincode by label position
            if idx < 2 and values[1]:
                # Determine run range for text after the word 'pincode'/'pin code'
                txt = p.text
                low = txt.lower()
                pin_label_pos = -1
                for key in ("pincode", "pin code"):
                    pin_label_pos = low.find(key)
                    if pin_label_pos != -1:
                        # move to after possible ':' or '-' and space
                        j = pin_label_pos + len(key)
                        while j < len(txt) and txt[j] in [':', '-', ' ', '\u00A0']:
                            j += 1
                        pin_label_pos = j
                        break
                # Walk runs and set the first run starting at/after pincode label position if it's empty/placeholder
                if pin_label_pos != -1:
                    pos = 0
                    candidate_after = None
                    candidate_after_highlight = None
                    for r in p.runs:
                        begin = pos
                        end = pos + len(r.text)
                        pos = end
                        if begin < pin_label_pos:
                            continue
                        # first run after the label
                        if candidate_after is None:
                            candidate_after = r
                        try:
                            if r.font.highlight_color == WD_COLOR_INDEX.YELLOW and candidate_after_highlight is None:
                                candidate_after_highlight = r
                        except Exception:
                            pass
                    target_run = candidate_after_highlight or candidate_after
                    if target_run is not None:
                        target_run.text = str(values[1])
                    else:
                        # As a last resort, append a run
                        try:
                            p.add_run(str(values[1]))
                        except Exception:
                            pass
            # continue with other labels too (in case paragraph also contains others)
        # Do scoped replacement for each label
        for label, value in targets.items():
            replace_in_paragraph(p, label, value, label_list)

    # Final cleanup: remove any leftover demo placeholders like 'replace ... here'
    for p in search_paras:
        for r in p.runs:
            try:
                if r.font.highlight_color == WD_COLOR_INDEX.YELLOW and r.text.strip().lower().startswith("replace"):
                    r.text = ""
            except Exception:
                pass


def clear_all_highlights(doc: DocxDocument) -> None:
    """Remove highlight formatting from all runs in the document (paragraphs and tables)."""
    for p in iter_paragraphs_and_cells(doc):
        for r in p.runs:
            try:
                # Setting to None clears any highlight color
                r.font.highlight_color = None
            except Exception:
                pass

# ---------------------------
# Core required functions
# ---------------------------

def _template_for_product(product: str) -> str:
    """Return the template path based on selected product."""
    pl = (product or "").lower()
    if "5.5" in pl:
        return TEMPLATE_PATH_55
    # default to 3.3 template
    return TEMPLATE_PATH

def create_invoice(form_data: Dict[str, str], template_path: str) -> Tuple[Optional[str], Optional[str]]:
    """Create invoice: generate DOCX (temporary), convert to PDF, save DB record.
    Returns (None, pdf_path or None if conversion failed). We no longer persist DOCX files.
    """
    ensure_dirs()
    conn = get_conn()
    try:
        qno = next_quotation_no(conn)
        form_data = dict(form_data)
        form_data["quotation_no"] = qno

        # Prepare values order for highlighted replacements
        # Order assumption based on the template sample provided:
        # [customer_name, location, city, state, pincode, mobile, product, quotation_no, date_of_quotation, validity_date]
        values = [
            form_data.get("customer_name", ""),
            form_data.get("location", ""),
            form_data.get("city", ""),
            form_data.get("state", ""),
            form_data.get("pincode", ""),
            form_data.get("mobile", ""),
            form_data.get("product", ""),
            form_data.get("quotation_no", ""),
            form_data.get("date_of_quotation", ""),
            form_data.get("validity_date", ""),
        ]

        docx_bytes = generate_docx(values, form_data, template_path)
        # Save DOCX temporarily (needed for conversion) – will delete after PDF is created
        # Use only Quotation No for file naming so one quotation -> one PDF file consistently
        safe_qno = safe_filename(form_data["quotation_no"]) if form_data.get("quotation_no") else "qno"
        base_name = f"{safe_qno}"
        docx_path = os.path.join(DOCX_DIR, f"{base_name}.docx")
        # Safe overwrite if file exists
        try:
            if os.path.exists(docx_path):
                os.remove(docx_path)
        except Exception:
            pass
        with open(docx_path, "wb") as f:
            f.write(docx_bytes.getvalue())

        # Convert to PDF
        target_pdf = os.path.join(PDF_DIR, f"{base_name}.pdf")
        try:
            if os.path.exists(target_pdf):
                os.remove(target_pdf)
        except Exception:
            pass
        pdf_path = convert_to_pdf(docx_path, target_pdf)

        # Always delete temporary DOCX (do not persist word files)
        try:
            if os.path.exists(docx_path):
                os.remove(docx_path)
        except Exception:
            pass
        persisted_docx_path: Optional[str] = None

        # Save DB
        save_to_db(conn, {
            **form_data,
            "docx_path": persisted_docx_path,
            "pdf_path": pdf_path,
        })
        return persisted_docx_path, pdf_path
    finally:
        conn.close()


def _render_pdf_preview(pdf_path: str, height: int = 700) -> None:
    """Render a PDF inline using PDF.js to avoid Chrome blocking the built-in viewer in sandboxed iframes."""
    try:
        with open(pdf_path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode("utf-8")
        # Minimal PDF.js renderer for all pages
        html = f"""
        <div id="pdf_root"></div>
        <script src="https://unpkg.com/pdfjs-dist@3.11.174/build/pdf.min.js"></script>
        <script>
        (function() {{
            const pdfData = atob('{b64}');
            const bytes = new Uint8Array(pdfData.length);
            for (let i = 0; i < pdfData.length; i++) bytes[i] = pdfData.charCodeAt(i);
            const CMAP_URL = 'https://unpkg.com/pdfjs-dist@3.11.174/cmaps/';
            const ROOT = document.getElementById('pdf_root');
            ROOT.style.border = '1px solid #e5e7eb';
            ROOT.style.borderRadius = '10px';
            ROOT.style.padding = '8px';
            const loadingTask = window['pdfjsLib'].getDocument({{ data: bytes, cMapUrl: CMAP_URL, cMapPacked: true }});
            loadingTask.promise.then(function(pdf) {{
                const scale = 1.1;
                const renderPage = function(num) {{
                    pdf.getPage(num).then(function(page) {{
                        const viewport = page.getViewport({{ scale }});
                        const canvas = document.createElement('canvas');
                        canvas.style.display = 'block';
                        canvas.style.margin = '0 auto 8px auto';
                        const context = canvas.getContext('2d');
                        canvas.height = viewport.height;
                        canvas.width = viewport.width;
                        ROOT.appendChild(canvas);
                        page.render({{ canvasContext: context, viewport: viewport }});
                    }});
                }};
                for (let i = 1; i <= pdf.numPages; i++) renderPage(i);
            }}).catch(function(err) {{
                ROOT.innerHTML = '<div style="color:#ef4444">Failed to load preview.</div>';
                console.error(err);
            }});
        }})();
        </script>
        """
        components.html(html, height=height, scrolling=True)
    except Exception:
        st.warning("Preview not available.")


def generate_docx(values_in_order: List[str], form_data: Dict[str, str], template_path: str) -> io.BytesIO:
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found at {template_path}")
    doc = Document(template_path)

    # Replace values by labels for accuracy
    replace_by_labels(doc, form_data)

    # Normalize layout to minimize LO vs Word differences
    normalize_layout(doc)

    # Remove any yellow highlighting so final PDF has clean text
    clear_all_highlights(doc)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def normalize_layout(doc: DocxDocument) -> None:
    """Stabilize table layout and currency formatting so LibreOffice doesn't wrap unexpectedly.
    - Turn off table autofit so column widths are respected
    - Apply fixed widths to the items table columns when detected
    - Right-align numeric/currency columns
    - Replace "₹ " with non-breaking space variant "₹\u00A0" so symbol sticks to the amount
    - Keep two-column detail table widths reasonable to avoid label/value wrapping
    """
    # Helper to set entire column width
    def set_col_width(table: Table, col_idx: int, width_in: float):
        w = Inches(width_in)
        for r in table.rows:
            try:
                r.cells[col_idx].width = w
            except Exception:
                pass

    # Replace rupee+space globally to prevent breaks
    for p in iter_paragraphs_and_cells(doc):
        for r in p.runs:
            try:
                if "₹ " in r.text:
                    r.text = r.text.replace("₹ ", "₹\u00A0")
            except Exception:
                pass

    for table in doc.tables:
        try:
            table.autofit = False
        except Exception:
            pass

        # Try to detect the items table by headers
        headers = []
        try:
            if table.rows:
                headers = [c.text.strip().lower() for c in table.rows[0].cells]
        except Exception:
            headers = []

        if headers and ("item name" in headers or "itemname" in headers) and ("amount" in headers):
            # Approximate column widths in inches matching an A4 portrait printable width (~6.2in content area)
            # [S.No, Item name, Qty, Price/Unit, GST(%), GST(Amount), Amount]
            col_widths = [0.5, 3.0, 0.7, 1.0, 0.8, 1.1, 1.1]
            for idx, w in enumerate(col_widths):
                if idx < len(table.columns):
                    set_col_width(table, idx, w)

            # Right-align numeric columns
            num_cols = []
            for key in ("price/ unit", "price/unit", "gst (amount)", "gst amount", "amount", "gst (%)", "gst%"):
                if key in headers:
                    num_cols.append(headers.index(key))
            # Fallback known positions if headers not matched precisely
            if not num_cols and len(headers) >= 7:
                num_cols = [3, 4, 5, 6]
            for r in table.rows:
                for ci in num_cols:
                    if ci < len(r.cells):
                        for p in r.cells[ci].paragraphs:
                            try:
                                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            except Exception:
                                pass
                        for p in r.cells[ci].paragraphs:
                            for run in p.runs:
                                try:
                                    if "₹ " in run.text:
                                        run.text = run.text.replace("₹ ", "₹\u00A0")
                                except Exception:
                                    pass
        else:
            # Heuristic for 2-column details table (labels/values block)
            try:
                if len(table.columns) == 2 and len(table.rows) >= 2:
                    # Allocate ~62% / 38% of width
                    set_col_width(table, 0, 3.8)
                    set_col_width(table, 1, 2.3)
            except Exception:
                pass


def convert_to_pdf(docx_path: str, target_pdf_path: str) -> Optional[str]:
    # Ensure target directory exists
    try:
        os.makedirs(os.path.dirname(target_pdf_path), exist_ok=True)
    except Exception:
        pass

    # 1) Try LibreOffice (fast, headless) if available (works on Linux/Streamlit Cloud and Windows if installed)
    try:
        # Try PATH first
        soffice = shutil.which("soffice") or shutil.which("soffice.exe")
        # Allow overriding via environment variable
        if not soffice:
            env_lo = os.environ.get("LIBREOFFICE_PATH")
            if env_lo and os.path.exists(env_lo):
                soffice = env_lo
        # Try common Windows install path
        if not soffice:
            win_lo = r"C:\\Program Files\\LibreOffice\\program\\soffice.exe"
            if os.path.exists(win_lo):
                soffice = win_lo
        if soffice and os.path.exists(docx_path):
            outdir = os.path.dirname(target_pdf_path)
            cmd = [
                soffice,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                outdir,
                docx_path,
            ]
            subprocess.run(cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=60)
            base = os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
            produced = os.path.join(outdir, base)
            if os.path.exists(produced):
                if os.path.abspath(produced) != os.path.abspath(target_pdf_path):
                    try:
                        if os.path.exists(target_pdf_path):
                            os.remove(target_pdf_path)
                    except Exception:
                        pass
                    os.replace(produced, target_pdf_path)
                return target_pdf_path
    except Exception:
        pass

    # 2) Fallback: Word via docx2pdf (Windows only)
    try:
        src = os.path.abspath(docx_path)
        dst = os.path.abspath(target_pdf_path)
        # Try file-to-file
        docx2pdf_convert(src, dst)
        if os.path.exists(dst):
            return dst
        # Try file-to-directory (docx2pdf will name the PDF same as DOCX base)
        outdir = os.path.dirname(dst)
        os.makedirs(outdir, exist_ok=True)
        docx2pdf_convert(src, outdir)
        produced = os.path.join(outdir, os.path.splitext(os.path.basename(src))[0] + ".pdf")
        if os.path.exists(produced):
            # Move/rename to target path if needed
            if os.path.abspath(produced) != os.path.abspath(dst):
                try:
                    if os.path.exists(dst):
                        os.remove(dst)
                except Exception:
                    pass
                os.replace(produced, dst)
            return dst
    except Exception:
        pass
    return None


def save_to_db(conn: sqlite3.Connection, record: Dict[str, str]) -> None:
    """Save a record to the database."""
    now = datetime.now().isoformat(timespec="seconds")
    conn.execute(
        """
        INSERT INTO invoices (
            quotation_no, product, customer_name, mobile, location, city, state, pincode,
            staff_name, date_of_quotation, validity_date, docx_path, pdf_path, created_at, updated_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            record.get("quotation_no"),
            record.get("product"),
            record.get("customer_name"),
            record.get("mobile"),
            record.get("location"),
            record.get("city"),
            record.get("state"),
            record.get("pincode"),
            record.get("staff_name"),
            record.get("date_of_quotation"),
            record.get("validity_date"),
            record.get("docx_path"),
            record.get("pdf_path"),
            now,
            now,
        ),
    )
    conn.commit()


def load_invoices() -> pd.DataFrame:
    """Load invoices from the database."""
    conn = get_conn()
    try:
        df = pd.read_sql_query(
            "SELECT id, customer_name, mobile, product, date_of_quotation, quotation_no, docx_path, pdf_path FROM invoices ORDER BY id DESC",
            conn,
        )
    finally:
        conn.close()
    return df


def delete_invoice(inv_id: int) -> None:
    """Delete an invoice from the database."""
    conn = get_conn()
    try:
        cur = conn.cursor()
        cur.execute("SELECT docx_path, pdf_path FROM invoices WHERE id = ?", (inv_id,))
        row = cur.fetchone()
        if row:
            docx_path, pdf_path = row
            try:
                if docx_path and os.path.exists(docx_path):
                    os.remove(docx_path)
            except Exception:
                pass
            try:
                if pdf_path and os.path.exists(pdf_path):
                    os.remove(pdf_path)
            except Exception:
                pass
        conn.execute("DELETE FROM invoices WHERE id = ?", (inv_id,))
        conn.commit()
    finally:
        conn.close()


def edit_agreement(agr_id: int, tags: Dict[str, str]) -> Optional[str]:
    """Regenerate an agreement PDF for the given record, keeping the same agreement_no and
    replacing the previous PDF in-place. Also updates name/number/address/date fields in DB.
    Returns the new PDF path (or None if conversion failed)."""
    rec = fetch_agreement_record(agr_id)
    if not rec:
        raise ValueError("Agreement not found")
    ensure_dirs()
    base = safe_filename(rec.get("agreement_no", f"AGR-{agr_id}"))
    template_path = os.path.join(os.getcwd(), "templates", "agreement template.docx")
    if not os.path.exists(template_path):
        raise FileNotFoundError("templates/agreement template.docx not found")

    temp_docx = os.path.join(DOCX_DIR, f"{base}.docx")
    try:
        if os.path.exists(temp_docx):
            os.remove(temp_docx)
    except Exception:
        pass

    # Populate using XML-level tag replacement
    _docx_zip_replace_tags(template_path, temp_docx, {
        "Date": tags.get("Date", rec.get("date", "")),
        "Name": tags.get("Name", rec.get("name", "")),
        "Address": tags.get("Address", rec.get("address", "")),
        "Number": tags.get("Number", rec.get("number", "")),
    })
    # Additional pass for split runs
    try:
        doc = Document(temp_docx)
        _replace_tags_in_docx(doc, {
            "Date": tags.get("Date", rec.get("date", "")),
            "Name": tags.get("Name", rec.get("name", "")),
            "Address": tags.get("Address", rec.get("address", "")),
            "Number": tags.get("Number", rec.get("number", "")),
        })
        clear_all_highlights(doc)
        # remove character spacing
        try:
            NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            for p in iter_paragraphs_and_cells(doc):
                for r in p.runs:
                    try:
                        rpr = r._element.rPr
                        if rpr is not None:
                            for node in r._element.xpath('.//w:spacing|.//w:kern', namespaces=NS):
                                parent = node.getparent()
                                if parent is not None:
                                    parent.remove(node)
                    except Exception:
                        pass
        except Exception:
            pass
        doc.save(temp_docx)
    except Exception:
        pass

    target_pdf = os.path.join(AGREEMENT_PDF_DIR, f"{base}.pdf")
    try:
        if os.path.exists(target_pdf):
            os.remove(target_pdf)
    except Exception:
        pass
    pdf_path = _convert_to_pdf_word_first(temp_docx, target_pdf)
    if pdf_path:
        try:
            if os.path.exists(temp_docx):
                os.remove(temp_docx)
        except Exception:
            pass

    conn = get_conn()
    try:
        conn.execute(
            """
            UPDATE agreements
            SET name=?, number=?, address=?, date=?, agreement_pdf_path=?
            WHERE id=?
            """,
            (
                tags.get("Name", rec.get("name", "")),
                tags.get("Number", rec.get("number", "")),
                tags.get("Address", rec.get("address", "")),
                tags.get("Date", rec.get("date", "")),
                pdf_path,
                agr_id,
            ),
        )
        conn.commit()
    finally:
        conn.close()
    return pdf_path


def edit_invoice(inv_id: int, form_data: Dict[str, str], template_path: str) -> Tuple[Optional[str], Optional[str]]:
    """Update record and regenerate files. Returns (None, pdf_path).
    We no longer persist DOCX files; use temporary DOCX for conversion only.
    """
    ensure_dirs()
    conn = get_conn()
    try:
        # Fetch existing quotation_no and existing file paths to keep it stable and replace old PDF
        cur = conn.cursor()
        cur.execute("SELECT quotation_no, docx_path, pdf_path FROM invoices WHERE id = ?", (inv_id,))
        row = cur.fetchone()
        if not row:
            raise ValueError("Invoice not found")
        quotation_no, old_docx_path, old_pdf_path = row[0], row[1], row[2]
        form_data = dict(form_data)
        form_data["quotation_no"] = quotation_no

        values = [
            form_data.get("customer_name", ""),
            form_data.get("location", ""),
            form_data.get("city", ""),
            form_data.get("state", ""),
            form_data.get("pincode", ""),
            form_data.get("mobile", ""),
            form_data.get("product", ""),
            form_data.get("quotation_no", ""),
            form_data.get("date_of_quotation", ""),
            form_data.get("validity_date", ""),
        ]
        docx_bytes = generate_docx(values, form_data, template_path)

        # Use only Quotation No for file naming so one quotation -> one PDF file
        safe_qno = safe_filename(form_data["quotation_no"]) if form_data.get("quotation_no") else "qno"
        base_name = f"{safe_qno}"
        temp_docx_path = os.path.join(DOCX_DIR, f"{base_name}.docx")
        with open(temp_docx_path, "wb") as f:
            f.write(docx_bytes.getvalue())

        pdf_path = convert_to_pdf(temp_docx_path, os.path.join(PDF_DIR, f"{base_name}.pdf"))

        # Always delete temporary DOCX and remove any previously stored DOCX
        try:
            if os.path.exists(temp_docx_path):
                os.remove(temp_docx_path)
        except Exception:
            pass
        try:
            if old_docx_path and os.path.exists(old_docx_path):
                os.remove(old_docx_path)
        except Exception:
            pass

        # If previous PDF exists and path differs from new target, delete it
        target_pdf_path = os.path.join(PDF_DIR, f"{base_name}.pdf")
        try:
            if old_pdf_path and os.path.exists(old_pdf_path) and os.path.abspath(old_pdf_path) != os.path.abspath(target_pdf_path):
                os.remove(old_pdf_path)
        except Exception:
            pass

        now = datetime.now().isoformat(timespec="seconds")
        conn.execute(
            """
            UPDATE invoices SET
                product=?, customer_name=?, mobile=?, location=?, city=?, state=?, pincode=?, staff_name=?,
                date_of_quotation=?, validity_date=?, docx_path=?, pdf_path=?, updated_at=?
            WHERE id=?
            """,
            (
                form_data.get("product"),
                form_data.get("customer_name"),
                form_data.get("mobile"),
                form_data.get("location"),
                form_data.get("city"),
                form_data.get("state"),
                form_data.get("pincode"),
                form_data.get("staff_name"),
                form_data.get("date_of_quotation"),
                form_data.get("validity_date"),
                None,
                pdf_path,
                now,
                inv_id,
            ),
        )
        conn.commit()
        return None, pdf_path
    finally:
        conn.close()


# ---------------------------
# Agreements: parsing, storage, generation
# ---------------------------

AGREEMENT_PREFIX = "AGR"

def _next_agreement_no(conn: sqlite3.Connection) -> str:
    """Generate next agreement number like AGR-YYYYMM-XXXX."""
    now = datetime.now()
    ym = now.strftime("%Y%m")
    cur = conn.cursor()
    cur.execute("SELECT COUNT(*) FROM agreements WHERE agreement_no LIKE ?", (f"{AGREEMENT_PREFIX}-{ym}-%",))
    n = cur.fetchone()[0] or 0
    return f"{AGREEMENT_PREFIX}-{ym}-{n+1:04d}"


def _save_feasibility_pdf(file_bytes: bytes, filename_hint: str) -> str:
    """Save feasibility PDF to a single canonical file to avoid duplicates across reruns.
    We always overwrite FEASIBILITY_DIR/feasibility.pdf
    """
    ensure_dirs()
    path = os.path.join(FEASIBILITY_DIR, "feasibility.pdf")
    try:
        if os.path.exists(path):
            os.remove(path)
    except Exception:
        pass
    with open(path, "wb") as f:
        f.write(file_bytes)
    # Log the upload event in DB
    conn = get_conn()
    try:
        conn.execute("INSERT INTO feasibility_events (uploaded_at) VALUES (?)", (datetime.now().isoformat(timespec="seconds"),))
        conn.commit()
    finally:
        conn.close()
    return path


def _read_pdf_text(pdf_bytes: bytes) -> str:
    """Extract text, preferring pdfminer for higher fidelity, with PyPDF2 fallback."""
    # 1) Try pdfminer.six (best for structured text in tables)
    try:
        from pdfminer.high_level import extract_text as _pdfminer_extract
        bio1 = io.BytesIO(pdf_bytes)
        txt = _pdfminer_extract(bio1) or ""
        if txt and txt.strip():
            return txt
    except Exception:
        pass
    # 2) Fallback to PyPDF2
    try:
        from PyPDF2 import PdfReader
        bio2 = io.BytesIO(pdf_bytes)
        reader = PdfReader(bio2)
        texts = []
        for page in reader.pages:
            try:
                texts.append(page.extract_text() or "")
            except Exception:
                pass
        return "\n".join(texts)
    except Exception:
        return ""


def _extract_fields_from_text(text: str) -> Dict[str, str]:
    """Extract Date, Name of Applicant, Mobile No, Address of Premises for Installation.
    Return keys: Date, Name, Address, Number.
    """
    t = re.sub(r"\r", "\n", text or "")
    t = re.sub(r"\u00A0", " ", t)
    lines = [ln.strip() for ln in t.splitlines() if ln.strip()]
    joined = "\n".join(lines)

    def value_after_inline(ln: str) -> Optional[str]:
        # Split by common separators
        for sep in [":", "-", "\t", "  "]:
            if sep in ln:
                parts = ln.split(sep, 1)
                if len(parts) == 2:
                    return parts[1].strip()
        # If no separator, take last token fallback (not ideal)
        toks = ln.split()
        if len(toks) > 1:
            return " ".join(toks[1:]).strip()
        return None

    def find_after(label_patterns: List[str]) -> Optional[str]:
        for pat in label_patterns:
            m = re.search(rf"{pat}\s*[:\-]\s*(.+)", joined, flags=re.IGNORECASE)
            if m:
                val = m.group(1).strip()
                val = re.split(r"\s{2,}|\s*(?:Name of Applicant|Mobile\s*No|Address of Premises for Installation|Date)\s*[:\-]", val, maxsplit=1, flags=re.IGNORECASE)[0].strip()
                return val
        for idx, ln in enumerate(lines):
            for pat in label_patterns:
                if re.search(rf"^{pat}\s*[:\-]?\s*$", ln, flags=re.IGNORECASE):
                    if idx + 1 < len(lines):
                        return lines[idx + 1].strip()
        return None

    # Primary label-based capture
    date_val = find_after([r"Date"]) or None
    name_val = find_after([r"Name\s+of\s+Applicant", r"Applicant\s*Name"]) or None
    mobile_val = find_after([r"Mobile\s*No", r"Mobile\s*Number", r"Contact\s*No"]) or None
    # Address block: capture block after "Address of Premises for installation" up to next numbered section or blank
    addr_idx = None
    for i, ln in enumerate(lines):
        if re.search(r"Address\s+of\s+Premises\s+for\s+installation", ln, flags=re.IGNORECASE):
            addr_idx = i
            break
    address_val = None
    if addr_idx is not None:
        block = []
        for j in range(addr_idx + 1, min(addr_idx + 10, len(lines))):
            l = lines[j]
            if re.match(r"^\d+\.|^Feasibility\s+Approval\s+Details", l, flags=re.IGNORECASE):
                break
            if re.match(r"^(From|To)\b", l, flags=re.IGNORECASE):
                break
            block.append(l)
        # Extract parts
        base_addr = None
        district = None
        state = None
        pincode = None
        for b in block:
            if re.search(r"^District\s*:\s*", b, flags=re.IGNORECASE):
                district = re.sub(r"^District\s*:\s*", "", b, flags=re.IGNORECASE).strip().strip(',')
            elif re.search(r"^State\s*:\s*", b, flags=re.IGNORECASE):
                state = re.sub(r"^State\s*:\s*", "", b, flags=re.IGNORECASE).strip().strip(',')
            elif re.search(r"^(PIN\s*Code|Pincode)\s*:\s*", b, flags=re.IGNORECASE):
                pincode = re.sub(r"^(PIN\s*Code|Pincode)\s*:\s*", "", b, flags=re.IGNORECASE).strip().strip(',')
            elif not base_addr:
                base_addr = b.strip().strip(',')
        if base_addr or district or state or pincode:
            # Compose without labels, e.g., "12/21B, Ashoke Nagar Road, Barasat, West Bengal, 112233"
            parts = [x for x in [base_addr, district, state, pincode] if x]
            addr = ", ".join(parts)
            # Normalize spacing: collapse multiple spaces; trim spaces around commas
            addr = re.sub(r"\s+", " ", addr)
            addr = re.sub(r"\s*,\s*", ", ", addr)
            addr = re.sub(r",\s*,+", ", ", addr)
            address_val = addr.strip()

    # Fallbacks if primary capture failed
    joined_lower = joined.lower()
    if not mobile_val:
        # Row-based scan for Mobile No – strictly target the same row then immediate next few lines
        mob_row_pattern = re.compile(r"\b(?:mobile\s*(?:no\.?|number)|contact\s*no\.?|phone\s*no\.?)\b", re.IGNORECASE)
        for i, ln in enumerate(lines):
            if mob_row_pattern.search(ln):
                # 1) Try to take digits from the same line after the label
                tail = mob_row_pattern.split(ln, maxsplit=1)[-1]
                # remove separators then extract digits
                d_same = re.findall(r"\d{8,13}", tail)
                if not d_same:
                    # Also handle formats with spaces/dashes in the number
                    d_same2 = re.findall(r"(?:\d[\s\-]?){9,14}\d", tail)
                    if d_same2:
                        d_join = re.sub(r"\D", "", d_same2[-1])
                        if 8 <= len(d_join) <= 13:
                            d_same = [d_join]
                if d_same:
                    # prefer 10-digit
                    pick = [x for x in d_same if len(x) == 10]
                    mobile_val = (pick[-1] if pick else d_same[-1])
                    break
                # 2) Look at the next few lines for the value cell; pick first that has an 8–13 digit sequence
                grabbed = False
                for j in range(i+1, min(i+5, len(lines))):
                    cand = lines[j].strip()
                    if not cand:
                        continue
                    # skip if looks like another label row
                    if re.search(r":|email|consumer|category|address|application|reference|discom|particulars|details", cand, flags=re.IGNORECASE):
                        continue
                    d_next = re.findall(r"\d{8,13}", cand)
                    if not d_next:
                        d_next2 = re.findall(r"(?:\d[\s\-]?){9,14}\d", cand)
                        if d_next2:
                            d_next = [re.sub(r"\D", "", d_next2[-1])]
                    if d_next:
                        pick = [x for x in d_next if len(x) == 10]
                        mobile_val = (pick[-1] if pick else d_next[-1])
                        grabbed = True
                        break
                if grabbed:
                    break
    # If still not found, pick best digit-only candidate of length 10-13
    if not mobile_val:
        cands = re.findall(r"\b\d{10,13}\b", joined)
        if cands:
            # prefer 10-digit, else longest
            ten = [c for c in cands if len(c) == 10]
            mobile_val = ten[0] if ten else max(cands, key=len)
    if not date_val:
        # Header style: granted on date: dd-mm-yyyy
        m = re.search(r"granted\s+on\s+date\s*[:\-]?\s*([0-9]{1,2}[^0-9A-Za-z][0-9]{1,2}[^0-9A-Za-z][0-9]{2,4})", joined_lower, flags=re.IGNORECASE)
        if not m:
            m = re.search(r"date\s*[:\-]?\s*([0-9]{1,2}[^0-9A-Za-z][0-9]{1,2}[^0-9A-Za-z][0-9]{2,4})", joined_lower, flags=re.IGNORECASE)
        if m:
            date_val = m.group(1)
    def is_valid_name(s: Optional[str]) -> bool:
        if not s:
            return False
        t = s.strip()
        if not t:
            return False
        low = t.lower()
        # disqualify obvious labels
        forbidden = [
            'mobile', 'phone', 'email', 'email id', 'consumer', 'category', 'address',
            'application', 'reference', 'number', 'no.', 'capacity', 'kwp', 'kw', 'discom',
        ]
        if any(k in low for k in forbidden):
            return False
        # must contain letters and not be mostly digits
        letters = len(re.findall(r"[A-Za-z]", t))
        digits = len(re.findall(r"\d", t))
        return letters >= 2 and letters > digits

    if not is_valid_name(name_val):
        # Row-based scan (handles tables) around the label row
        for i, ln in enumerate(lines):
            if re.search(r"\bname\s*of\s*applicant\b", ln, flags=re.IGNORECASE):
                # prefer taking the RIGHT column/value on the same line
                # patterns: <label> : <value>  OR  <label> <many spaces or tab> <value>
                m_same = re.search(r"\bname\s*of\s*applicant\b\s*(?:[:\-]|\s{2,}|\t)\s*(.+)$", ln, flags=re.IGNORECASE)
                v = None
                if m_same:
                    v = m_same.group(1).strip()
                else:
                    # As a fallback, split by two+ spaces and take last chunk to mimic table column split
                    parts = re.split(r"\s{2,}|\t", ln)
                    # ensure left-most chunk contains the label; take last non-empty chunk as value
                    if len(parts) >= 2 and re.search(r"\bname\s*of\s*applicant\b", parts[0], flags=re.IGNORECASE):
                        v = parts[-1].strip()
                if not is_valid_name(v):
                    # try the very next non-empty line as the details cell
                    # but skip if it obviously looks like another label row (contains ':' or known keywords)
                    for j in range(i+1, min(i+5, len(lines))):
                        cand = lines[j].strip()
                        lowc = cand.lower()
                        if not cand:
                            continue
                        if ':' in cand:
                            continue
                        if any(k in lowc for k in ['mobile', 'phone', 'email', 'consumer', 'category', 'address', 'application', 'reference']):
                            continue
                        if is_valid_name(cand):
                            v = cand
                            break
                if not is_valid_name(v):
                    # search next few rows for a name-like candidate
                    for j in range(i+1, min(i+8, len(lines))):
                        cand = lines[j].strip()
                        if is_valid_name(cand):
                            v = cand
                            break
                    # also look a couple of lines above in case of wrap
                    if not is_valid_name(v):
                        for j in range(max(0, i-3), i):
                            cand = lines[j].strip()
                            if is_valid_name(cand):
                                v = cand
                                break
                if is_valid_name(v):
                    name_val = v
                    break
        # Header-style fallback like 'Shri/Smt <NAME> ...'
        if not is_valid_name(name_val):
            m = re.search(r"(Shri|Smt|Shri/Smt|Sh\.?/Smt\.?)\s*[:.-]?\s*([A-Z][A-Za-z .,-]+)", joined, flags=re.IGNORECASE)
            if m and is_valid_name(m.group(1)):
                name_val = (m.group(2) or m.group(1)).strip()

    if mobile_val:
        # Normalize: extract digits, prefer a clean 10-digit Indian-style number if present
        digits = re.findall(r"\d", mobile_val)
        if digits:
            s = "".join(digits)
            # If there is a 10-digit substring, take the last one
            m10 = re.findall(r"\d{10}", s)
            if m10:
                mobile_val = m10[-1]
            elif len(s) >= 10:
                mobile_val = s[-10:]
            else:
                mobile_val = s

    def norm_date(s: Optional[str]) -> Optional[str]:
        if not s:
            return None
        for fmt in ("%d/%m/%Y", "%d-%m-%Y", "%Y-%m-%d", "%d.%m.%Y", "%d %b %Y", "%d %B %Y"):
            try:
                d = datetime.strptime(s.strip(), fmt).date()
                return d.strftime("%d-%m-%Y")
            except Exception:
                continue
        m = re.search(r"(\d{1,2})[^0-9A-Za-z](\d{1,2})[^0-9A-Za-z](\d{2,4})", s)
        if m:
            d, mth, y = m.groups()
            y = y if len(y) == 4 else ("20" + y)
            try:
                return datetime(int(y), int(mth), int(d)).strftime("%d-%m-%Y")
            except Exception:
                pass
        return s.strip()

    return {
        "Date": norm_date(date_val) or "",
        "Name": (name_val or "").strip(),
        "Address": (address_val or "").strip(),
        "Number": (mobile_val or "").strip(),
    }


def _replace_tags_in_docx(doc: DocxDocument, mapping: Dict[str, str]) -> None:
    """Replace tag placeholders across the document.
    Strategy:
    1) First pass replaces sequences of YELLOW-highlighted runs whose combined text equals a tag like [Name].
       This preserves surrounding formatting and focuses only on highlighted placeholders.
    2) Second pass does a paragraph-level replace for any remaining tags (non-highlighted), by joining run texts,
       then writing the replaced string into the first run and clearing the rest. This is a fallback to catch
       placeholders that aren't highlighted, at the cost of losing mixed-run formatting inside that paragraph.
    """
    TAGS = {"[Date]": mapping.get("Date", ""),
            "[Name]": mapping.get("Name", ""),
            "[Address]": mapping.get("Address", ""),
            "[Number]": mapping.get("Number", "")}

    # 1) Replace in contiguous highlighted sequences
    for p in iter_paragraphs_and_cells(doc):
        runs = list(p.runs)
        i = 0
        while i < len(runs):
            # collect contiguous yellow runs
            j = i
            seq = []
            while j < len(runs):
                r = runs[j]
                try:
                    if r.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                        seq.append(r)
                        j += 1
                        continue
                except Exception:
                    pass
                break
            if seq:
                combined = "".join([x.text for x in seq])
                # exact tag match or contains tag (sometimes brackets split into separate runs)
                replaced = False
                for tag, val in TAGS.items():
                    if combined == tag or tag in combined:
                        # put value in first run, clear the rest
                        seq[0].text = val or ""
                        for k in range(1, len(seq)):
                            seq[k].text = ""
                        replaced = True
                        break
                i = j  # skip processed seq
                continue
            i += 1

    # 2) Fallback paragraph-wide substitution for any leftover naked tags
    for p in iter_paragraphs_and_cells(doc):
        try:
            full = "".join(r.text for r in p.runs)
            new_full = full
            for tag, val in TAGS.items():
                if tag in new_full:
                    new_full = new_full.replace(tag, val or "")
            if new_full != full and p.runs:
                p.runs[0].text = new_full
                for r in p.runs[1:]:
                    r.text = ""
        except Exception:
            pass


def _convert_to_pdf_word_first(docx_path: str, target_pdf_path: str) -> Optional[str]:
    """Prefer MS Word via docx2pdf for best layout fidelity on Windows, then fallback to LibreOffice."""
    # Try Word/docx2pdf first
    try:
        src = os.path.abspath(docx_path)
        dst = os.path.abspath(target_pdf_path)
        docx2pdf_convert(src, dst)
        if os.path.exists(dst):
            return dst
        outdir = os.path.dirname(dst)
        os.makedirs(outdir, exist_ok=True)
        docx2pdf_convert(src, outdir)
        produced = os.path.join(outdir, os.path.splitext(os.path.basename(src))[0] + ".pdf")
        if os.path.exists(produced):
            if os.path.abspath(produced) != os.path.abspath(dst):
                try:
                    if os.path.exists(dst):
                        os.remove(dst)
                except Exception:
                    pass
                os.replace(produced, dst)
            return dst
    except Exception:
        pass
    # Fallback to existing pipeline (LibreOffice then Word)
    return convert_to_pdf(docx_path, target_pdf_path)


def _docx_zip_replace_tags(template_path: str, out_path: str, mapping: Dict[str, str]) -> None:
    """Do a raw XML text replacement for placeholders across all word/*.xml parts to preserve layout.
    mapping keys should be simple text like 'Date', 'Name', etc., and placeholders are [Date], [Name], etc.
    """
    placeholders = {f"[{k}]": (v or "") for k, v in mapping.items()}
    with zipfile.ZipFile(template_path, 'r') as zin:
        with zipfile.ZipFile(out_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                    try:
                        xml = data.decode('utf-8')
                    except UnicodeDecodeError:
                        try:
                            xml = data.decode('cp1252')
                        except Exception:
                            xml = data.decode('utf-8', errors='ignore')
                    # Replace all placeholders (plain match and tag-insensitive regex match)
                    for ph, val in placeholders.items():
                        if ph in xml:
                            xml = xml.replace(ph, val)
                        # Tag-insensitive match: allow arbitrary XML tags and optional spaces inside brackets
                        # Example: matches "[Name]", "[ Name ]" even when split across <w:t> runs
                        def build_tag_insensitive_pattern(text: str) -> str:
                            core = text.strip("[]")
                            parts = [r"\[\s*"]
                            for ch in core:
                                parts.append(re.escape(ch))
                                parts.append(r"(?:\s*<[^>]+>\s*)*")
                            parts.append(r"\s*\]")
                            return "".join(parts)
                        pat = build_tag_insensitive_pattern(ph)
                        try:
                            xml = re.sub(pat, val, xml)
                        except Exception:
                            pass
                    # Strip character spacing and kerning tags that cause spaced-out letters
                    try:
                        # Remove self-closing and open/close variants
                        xml = re.sub(r"<w:spacing[^>]*/>", "", xml)
                        xml = re.sub(r"<w:spacing[^>]*>\s*</w:spacing>", "", xml)
                        xml = re.sub(r"<w:kern[^>]*/>", "", xml)
                        xml = re.sub(r"<w:kern[^>]*>\s*</w:kern>", "", xml)
                        # Replace distributed justification which creates per-character spacing
                        xml = re.sub(r"<w:jc[^>]*w:val=\"distribute\"[^>]*/>", "<w:jc w:val=\"left\"/>", xml)
                        xml = re.sub(r"<w:jc[^>]*w:val=\"thaiDistribute\"[^>]*/>", "<w:jc w:val=\"left\"/>", xml)
                        # Open/close jc tags as well
                        xml = re.sub(r"<w:jc[^>]*w:val=\"distribute\"[^>]*>\s*</w:jc>", "<w:jc w:val=\"left\"/>", xml)
                        xml = re.sub(r"<w:jc[^>]*w:val=\"thaiDistribute\"[^>]*>\s*</w:jc>", "<w:jc w:val=\"left\"/>", xml)
                        # Remove East-Asian spacing control that can expand characters
                        xml = re.sub(r"<w:characterSpacingControl[^>]*/>", "", xml)
                        xml = re.sub(r"<w:characterSpacingControl[^>]*>\s*</w:characterSpacingControl>", "", xml)
                        # DrawingML (text boxes) spacing/alignment
                        xml = re.sub(r"<a:spcPct[^>]*/>", "", xml)
                        xml = re.sub(r"<a:spcPts[^>]*/>", "", xml)
                        xml = re.sub(r"<a:spcPct[^>]*>\s*</a:spcPct>", "", xml)
                        xml = re.sub(r"<a:spcPts[^>]*>\s*</a:spcPts>", "", xml)
                        xml = re.sub(r"<a:algn[^>]*val=\"dist\"[^>]*/>", "<a:algn val=\"l\"/>", xml)
                        xml = re.sub(r"<a:algn[^>]*val=\"dist\"[^>]*>\s*</a:algn>", "<a:algn val=\"l\"/>", xml)
                    except Exception:
                        pass
                    data = xml.encode('utf-8')
                zout.writestr(item, data)


def generate_agreement_pdf(tags: Dict[str, str], feasibility_pdf_path: str) -> Tuple[Optional[str], Optional[str], str]:
    """Populate agreement template with tags at XML level and convert to PDF. Returns (docx_temp, pdf_path, agreement_no).

    Duplicate policy: If an agreement already exists with the same (Name, Number, Address),
    we raise DuplicateAgreementError so the UI can notify the user to check Generated Agreement.
    The user can then edit/replace from the Generated Agreement tab if needed.
    """
    ensure_dirs()
    # Duplicate detection
    name_v = (tags.get("Name", "") or "").strip()
    number_v = (tags.get("Number", "") or "").strip()
    address_v = (tags.get("Address", "") or "").strip()
    conn = get_conn()
    try:
        cur = conn.cursor()
        cur.execute(
            """
            SELECT id, agreement_no FROM agreements
            WHERE LOWER(COALESCE(name,'')) = LOWER(?)
              AND LOWER(COALESCE(address,'')) = LOWER(?)
              AND COALESCE(number,'') = ?
            LIMIT 1
            """,
            (name_v, address_v, number_v),
        )
        dup = cur.fetchone()
        if dup:
            raise DuplicateAgreementError(f"Duplicate exists with Agreement No: {dup[1]}")
    finally:
        conn.close()

    conn = get_conn()
    try:
        agreement_no = _next_agreement_no(conn)
    finally:
        conn.close()

    template_path = os.path.join(os.getcwd(), "templates", "agreement template.docx")
    if not os.path.exists(template_path):
        raise FileNotFoundError("templates/agreement template.docx not found")

    base = safe_filename(agreement_no)
    temp_docx = os.path.join(DOCX_DIR, f"{base}.docx")
    try:
        if os.path.exists(temp_docx):
            os.remove(temp_docx)
    except Exception:
        pass
    # XML-level replace preserves formatting and updates everywhere (including text boxes)
    _docx_zip_replace_tags(template_path, temp_docx, {
        "Date": tags.get("Date", ""),
        "Name": tags.get("Name", ""),
        "Address": tags.get("Address", ""),
        "Number": tags.get("Number", ""),
    })

    # Run python-docx pass too to catch tags that are split across runs within body paragraphs/tables
    try:
        doc = Document(temp_docx)
        # Helper to remove character spacing (tracking) so long addresses don't show spaced letters
        def _clear_char_spacing(_doc: DocxDocument) -> None:
            NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            # Clear on all runs
            for p in iter_paragraphs_and_cells(_doc):
                for r in p.runs:
                    try:
                        rpr = r._element.rPr
                        if rpr is not None:
                            for node in r._element.xpath('.//w:spacing|.//w:kern', namespaces=NS):
                                parent = node.getparent()
                                if parent is not None:
                                    parent.remove(node)
                    except Exception:
                        pass
            # Clear at style level (paragraph and character styles)
            try:
                for style in _doc.styles:
                    try:
                        el = style._element
                        for path in ['.//w:rPr/w:spacing', './/w:rPr/w:kern']:
                            for node in el.xpath(path, namespaces=NS):
                                parent = node.getparent()
                                if parent is not None:
                                    parent.remove(node)
                    except Exception:
                        continue
            except Exception:
                pass
        _replace_tags_in_docx(doc, {
            "Date": tags.get("Date", ""),
            "Name": tags.get("Name", ""),
            "Address": tags.get("Address", ""),
            "Number": tags.get("Number", ""),
        })
        clear_all_highlights(doc)
        _clear_char_spacing(doc)
        doc.save(temp_docx)
    except Exception:
        pass

    target_pdf = os.path.join(AGREEMENT_PDF_DIR, f"{base}.pdf")
    try:
        if os.path.exists(target_pdf):
            os.remove(target_pdf)
    except Exception:
        pass
    # Use Word-first for agreements to preserve template formatting
    pdf_path = _convert_to_pdf_word_first(temp_docx, target_pdf)

    # If conversion failed, keep DOCX so user can download; else delete temp DOCX
    if pdf_path:
        try:
            if os.path.exists(temp_docx):
                os.remove(temp_docx)
        except Exception:
            pass

    conn = get_conn()
    try:
        conn.execute(
            """
            INSERT INTO agreements (agreement_no, name, number, address, date, feasibility_pdf_path, agreement_pdf_path, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                agreement_no,
                tags.get("Name", ""),
                tags.get("Number", ""),
                tags.get("Address", ""),
                tags.get("Date", ""),
                feasibility_pdf_path,
                pdf_path,
                datetime.now().isoformat(timespec="seconds"),
            ),
        )
        conn.commit()
    finally:
        conn.close()

    return None, pdf_path, agreement_no


def load_agreements() -> pd.DataFrame:
    conn = get_conn()
    try:
        df = pd.read_sql_query(
            """
            SELECT id, agreement_no, name, number, address, date, feasibility_pdf_path, agreement_pdf_path, created_at
            FROM agreements
            ORDER BY id DESC
            """,
            conn,
        )
    finally:
        conn.close()
    return df


def fetch_agreement_record(agr_id: int) -> Optional[Dict[str, str]]:
    """Fetch a single agreement record by id as a dict."""
    conn = get_conn()
    try:
        cur = conn.cursor()
        cur.execute(
            """
            SELECT id, agreement_no, name, number, address, date, feasibility_pdf_path, agreement_pdf_path, created_at
            FROM agreements WHERE id=?
            """,
            (agr_id,),
        )
        r = cur.fetchone()
        if not r:
            return None
        return {
            "id": r[0],
            "agreement_no": r[1],
            "name": r[2],
            "number": r[3],
            "address": r[4],
            "date": r[5],
            "feasibility_pdf_path": r[6],
            "agreement_pdf_path": r[7],
            "created_at": r[8],
        }
    finally:
        conn.close()


def delete_agreement(agr_id: int) -> None:
    """Delete an agreement row and both associated PDFs (feasibility and agreement).
    This ensures a clean slate so re-creating does not hit UNIQUE/leftover-file issues.
    """
    conn = get_conn()
    try:
        cur = conn.cursor()
        cur.execute("SELECT feasibility_pdf_path, agreement_pdf_path FROM agreements WHERE id=?", (agr_id,))
        row = cur.fetchone()
        if row:
            feas, agrpdf = row
            # Delete both PDFs if they exist
            try:
                if agrpdf and os.path.exists(agrpdf):
                    os.remove(agrpdf)
            except Exception:
                pass
            try:
                if feas and os.path.exists(feas):
                    os.remove(feas)
            except Exception:
                pass
        conn.execute("DELETE FROM agreements WHERE id=?", (agr_id,))
        conn.commit()
    finally:
        conn.close()


def render_upload_feasibility_tab():
    st.subheader("Upload PMSG Approval of feasibility")
    st.caption("Upload only PDF. The file should match the template fields for accurate extraction.")
    up = st.file_uploader("Upload Feasibility PDF", type=["pdf"], accept_multiple_files=False)

    tags: Dict[str, str] = {}
    feasibility_pdf_path: Optional[str] = None
    if up is not None:
        pdf_bytes = up.read()
        feasibility_pdf_path = _save_feasibility_pdf(pdf_bytes, up.name)

        cprev, cform = st.columns([3, 2])
        with cprev:
            st.markdown("#### Preview: Uploaded Feasibility")
            try:
                _render_pdf_preview(feasibility_pdf_path, height=480)
            except Exception:
                st.info("Preview not available.")

        text = _read_pdf_text(pdf_bytes)
        tags = _extract_fields_from_text(text)

        with cform:
            st.markdown("#### Extracted Fields")
            st.markdown(f"**Name:** {tags.get('Name','')}")
            st.markdown(f"**Date:** {tags.get('Date','')}")
            st.markdown(f"**Phone:** {tags.get('Number','')}")
            st.text_area("Address:", value=(tags.get("Address", "") or ""), disabled=True, height=120)
            with st.expander("Debug: raw text (first 1200 chars)", expanded=False):
                st.code((text or "")[:1200] or "<empty>")
            st.markdown("---")
            if st.button("Create Agreement", use_container_width=True):
                # Use extracted values as-is (no manual editing per requirement)
                tags = {
                    "Date": tags.get("Date", ""),
                    "Name": tags.get("Name", ""),
                    "Address": tags.get("Address", ""),
                    "Number": tags.get("Number", ""),
                }
                try:
                    with st.spinner("Generating agreement…"):
                        _, agr_pdf, agr_no = generate_agreement_pdf(tags, feasibility_pdf_path)
                except DuplicateAgreementError:
                    st.error("Duplicate found - Check Generated invoices")
                    st.caption("An agreement with the same Name, Phone and Address already exists. Please review it in 'Generated Agreement'.")
                    return
                if agr_pdf and os.path.exists(agr_pdf):
                    st.success(f"Agreement generated: {agr_no}")
                    st.markdown("#### Preview: Agreement")
                    _render_pdf_preview(agr_pdf, height=480)
                    with open(agr_pdf, "rb") as f:
                        st.download_button(
                            "⬇️  Download Agreement PDF",
                            data=f.read(),
                            file_name=os.path.basename(agr_pdf),
                            mime="application/pdf",
                            use_container_width=True,
                        )
                    _render_mobile_share_button(agr_pdf, os.path.basename(agr_pdf))
                else:
                    st.error("Failed to generate agreement PDF. Ensure MS Word or LibreOffice is installed for DOCX→PDF conversion.")


def render_generated_agreements_tab():
    st.subheader("Generated Agreement")
    df = load_agreements()
    if df.empty:
        st.info("No agreements generated yet.")
        return

    # Filters
    c1, c2 = st.columns([2, 2])
    with c1:
        f_name = st.text_input("Filter by Customer Name", key="agreements_filter_name")
    with c2:
        f_id = st.text_input("Filter by Agreement No", key="agreements_filter_id")

    mask = pd.Series([True] * len(df))
    if f_name:
        mask &= df["name"].str.contains(f_name, case=False, na=False)
    if f_id:
        mask &= df["agreement_no"].astype(str).str.contains(f_id, case=False, na=False)
    filtered = df[mask].reset_index(drop=True)

    # Query param actions (preview/edit/delete)
    qp = st.query_params
    def _first(v):
        if v is None:
            return None
        if isinstance(v, list):
            return v[0] if v else None
        return v
    action = _first(qp.get("g_action"))
    action_id = _first(qp.get("g_id"))
    if action and action_id:
        try:
            rid = int(action_id)
            if action == "preview":
                st.session_state["agr_preview_id"] = rid
            elif action == "edit":
                st.session_state["agr_edit_id"] = rid
                st.session_state.pop("agr_preview_id", None)
            elif action == "delete":
                delete_agreement(rid)
                st.success("Deleted.")
            st.query_params.clear()
        except Exception:
            pass

    mobile_view = st.toggle("Mobile card view", value=True, key="agreements_mobile_card_toggle")

    # Light CSS similar to Search Invoice
    st.markdown(
        """
        <style>
        .card-title {font-weight:600; margin:0;}
        .meta {color:#6b7280; font-size:12px; margin: 0;}
        .stButton>button, .stDownloadButton>button {padding: 0.35rem 0.6rem; border-radius:999px; font-size:13px; min-width: 36px; height: 36px;}
        </style>
        """,
        unsafe_allow_html=True,
    )

    if mobile_view:
        for _, row in filtered.iterrows():
            rid = int(row["id"])
            pdf_ag = row.get("agreement_pdf_path") if "agreement_pdf_path" in row else None
            pdf_feas = row.get("feasibility_pdf_path") if "feasibility_pdf_path" in row else None
            with st.container(border=True):
                top_l, _ = st.columns([7, 3])
                with top_l:
                    st.markdown(f"<p class='card-title'>{row['name']}</p>", unsafe_allow_html=True)
                    st.markdown(f"<p class='meta'>Agreement No: {row['agreement_no']}</p>", unsafe_allow_html=True)

                # Row 1: Preview buttons
                p1, p2 = st.columns([1, 1])
                with p1:
                    if st.button("👁️  Preview Feasibility", key=f"ga_prev_f_{rid}", use_container_width=True, disabled=not (pdf_feas and os.path.exists(pdf_feas))):
                        st.session_state["agr_preview_feas_id"] = rid
                with p2:
                    if st.button("👁️  Preview Agreement", key=f"ga_prev_a_{rid}", use_container_width=True, disabled=not (pdf_ag and os.path.exists(pdf_ag))):
                        st.session_state["agr_preview_id"] = rid
                        st.session_state.pop("agr_preview_feas_id", None)

                # Row 2: Download Feasibility, Download Agreement, Edit, Share, Delete
                a1, a2, a3, a4, a5 = st.columns([1, 1, 1, 1, 1])
                with a1:
                    if pdf_feas and os.path.exists(pdf_feas):
                        with open(pdf_feas, "rb") as f:
                            st.download_button("⬇️  Download Feasibility", data=f.read(), file_name=os.path.basename(pdf_feas), mime="application/pdf", key=f"ga_dl_feas_{rid}", use_container_width=True)
                    else:
                        st.button("⬇️  Download Feasibility", disabled=True, key=f"ga_dl_feas_na_{rid}", use_container_width=True)
                with a2:
                    if pdf_ag and os.path.exists(pdf_ag):
                        with open(pdf_ag, "rb") as f:
                            st.download_button("⬇️  Download Agreement", data=f.read(), file_name=os.path.basename(pdf_ag), mime="application/pdf", key=f"ga_dl_ag_{rid}", use_container_width=True)
                    else:
                        st.button("⬇️  Download Agreement", disabled=True, key=f"ga_dl_ag_na_{rid}", use_container_width=True)
                with a3:
                    if st.button("✏️  Edit", key=f"ga_edit_{rid}", use_container_width=True):
                        st.session_state["agr_edit_id"] = rid
                        st.session_state.pop("agr_preview_id", None)
                        st.session_state.pop("agr_preview_feas_id", None)
                        st.rerun()
                with a4:
                    if pdf_ag and os.path.exists(pdf_ag):
                        _render_mobile_share_button(pdf_ag, os.path.basename(pdf_ag))
                    else:
                        st.button("Share PDF", disabled=True, key=f"ga_share_na_{rid}", use_container_width=True)
                with a5:
                    if st.button("🗑️  Delete", key=f"ga_del_{rid}", use_container_width=True):
                        # Clear state to avoid referencing a deleted record then rerun
                        st.session_state.pop("agr_preview_id", None)
                        st.session_state.pop("agr_preview_feas_id", None)
                        st.session_state.pop("agr_edit_id", None)
                        delete_agreement(rid)
                        st.success("Deleted.")
                        st.rerun()

                # Inline previews (full-width inside this card)
                if st.session_state.get("agr_preview_feas_id") == rid and pdf_feas and os.path.exists(pdf_feas):
                    _render_pdf_preview(pdf_feas, height=780)
                    if st.button("Close Feasibility preview", key=f"ga_close_prev_f_{rid}"):
                        st.session_state.pop("agr_preview_feas_id", None)
                        st.rerun()
                if st.session_state.get("agr_preview_id") == rid and pdf_ag and os.path.exists(pdf_ag):
                    _render_pdf_preview(pdf_ag, height=780)
                    if st.button("Close Agreement preview", key=f"ga_close_prev_{rid}"):
                        st.session_state.pop("agr_preview_id", None)
                        st.rerun()

                # Inline edit panel
                if st.session_state.get("agr_edit_id") == rid:
                    st.markdown("#### Edit Agreement")
                    rec = fetch_agreement_record(rid) or {}
                    c_cancel, _ = st.columns([1, 6])
                    with c_cancel:
                        if st.button("Close", key=f"ga_close_edit_{rid}"):
                            st.session_state.pop("agr_edit_id", None)
                            st.rerun()
                    with st.form(key=f"ga_edit_form_{rid}"):
                        name = st.text_input("Name", value=rec.get("name", ""))
                        number = st.text_input("Phone", value=rec.get("number", ""))
                        date = st.text_input("Date (DD-MM-YYYY)", value=rec.get("date", ""))
                        address = st.text_area("Address", value=rec.get("address", ""), height=100)
                        submitted = st.form_submit_button("Save & Regenerate")
                    if submitted:
                        new_tags = {"Name": name.strip(), "Number": number.strip(), "Date": date.strip(), "Address": address.strip()}
                        with st.spinner("Regenerating…"):
                            pdf = edit_agreement(rid, new_tags)
                        if pdf and os.path.exists(pdf):
                            st.success("Agreement updated.")
                            st.rerun()
        return

    # Desktop-like layout (header + rows)
    h1, h2, h3, h4, h5 = st.columns([2.5, 2.5, 2, 2, 2])
    with h1:
        st.markdown("**Customer**")
    with h2:
        st.markdown("**Agreement No**")
    with h3:
        st.markdown("**Date**")
    with h4:
        st.markdown("**Phone**")
    with h5:
        st.markdown("**Action**")

    for _, row in filtered.iterrows():
        rid = int(row["id"])
        pdf_ag = row.get("agreement_pdf_path") if "agreement_pdf_path" in row else None
        pdf_feas = row.get("feasibility_pdf_path") if "feasibility_pdf_path" in row else None
        c1, c2, c3, c4, c5 = st.columns([2.5, 2.5, 2, 2, 2])
        with c1:
            st.write(row["name"])  
        with c2:
            st.write(row["agreement_no"]) 
        with c3:
            st.write(row.get("date", "")) 
        with c4:
            st.write(row.get("number", "")) 
        with c5:
            a1, a2, a3, a4, a5 = st.columns([1, 1, 1, 1, 1])
            with a1:
                # Preview both (Agreement prioritized), then Feasibility below inline
                if st.button("👁️", key=f"ga_d_prev_{rid}", use_container_width=True):
                    st.session_state["agr_preview_id"] = rid
            with a2:
                if pdf_feas and os.path.exists(pdf_feas):
                    with open(pdf_feas, "rb") as f:
                        st.download_button("⬇️ F", f.read(), file_name=os.path.basename(pdf_feas), key=f"ga_d_dl_f_{rid}", use_container_width=True)
                else:
                    st.button("⬇️ F", disabled=True, key=f"ga_d_dl_f_na_{rid}", use_container_width=True)
            with a3:
                if pdf_ag and os.path.exists(pdf_ag):
                    with open(pdf_ag, "rb") as f:
                        st.download_button("⬇️ A", f.read(), file_name=os.path.basename(pdf_ag), key=f"ga_d_dl_a_{rid}", use_container_width=True)
                else:
                    st.button("⬇️ A", disabled=True, key=f"ga_d_dl_a_na_{rid}", use_container_width=True)
            with a4:
                if st.button("✏️", key=f"ga_d_edit_{rid}", use_container_width=True):
                    st.session_state["agr_edit_id"] = rid
                    st.session_state.pop("agr_preview_id", None)
                    st.rerun()
            with a5:
                if st.button("🗑️", key=f"ga_d_del_{rid}", use_container_width=True):
                    st.session_state.pop("agr_preview_id", None)
                    st.session_state.pop("agr_edit_id", None)
                    delete_agreement(rid)
                    st.success("Deleted.")
                    st.rerun()

        # Inline previews (full-width below row)
        if st.session_state.get("agr_preview_id") == rid:
            if pdf_feas and os.path.exists(pdf_feas):
                st.markdown("Feasibility Preview")
                _render_pdf_preview(pdf_feas, height=780)
            if pdf_ag and os.path.exists(pdf_ag):
                st.markdown("Agreement Preview")
                _render_pdf_preview(pdf_ag, height=780)
            if st.button("Close preview", key=f"ga_d_close_prev_{rid}"):
                st.session_state.pop("agr_preview_id", None)
                st.rerun()

        if st.session_state.get("agr_edit_id") == rid:
            st.markdown("#### Edit Agreement")
            rec = fetch_agreement_record(rid) or {}
            c_cancel, _ = st.columns([1, 6])
            with c_cancel:
                if st.button("Close", key=f"ga_d_close_edit_{rid}"):
                    st.session_state.pop("agr_edit_id", None)
                    st.rerun()
            with st.form(key=f"ga_d_edit_form_{rid}"):
                name = st.text_input("Name", value=rec.get("name", ""))
                number = st.text_input("Phone", value=rec.get("number", ""))
                date = st.text_input("Date (DD-MM-YYYY)", value=rec.get("date", ""))
                address = st.text_area("Address", value=rec.get("address", ""), height=100)
                submitted = st.form_submit_button("Save & Regenerate")
            if submitted:
                new_tags = {"Name": name.strip(), "Number": number.strip(), "Date": date.strip(), "Address": address.strip()}
                with st.spinner("Regenerating…"):
                    pdf = edit_agreement(rid, new_tags)
                if pdf and os.path.exists(pdf):
                    st.success("Agreement updated.")
                    st.rerun()
# ---------------------------
# Streamlit UI
# ---------------------------

def main():
    st.set_page_config(page_title=APP_TITLE, layout="wide")
    st.title(APP_TITLE)

    ensure_dirs()
    _ = get_conn()  # ensure DB exists

    # Sidebar navigation drawer
    st.sidebar.title("Navigation")
    section = st.sidebar.radio("Menu", ["Dashboard", "Invoice", "Agreement"], index=0)

    if section == "Dashboard":
        render_dashboard()

    elif section == "Invoice":
        tabs = st.tabs(["Create Invoice", "Search Invoice"])
        with tabs[0]:
            render_create_form(
                allowed_products=PRODUCT_OPTIONS,
                form_title="Create Invoice",
                key_ns="kwall",
            )
        with tabs[1]:
            render_search_tab()

    elif section == "Agreement":
        tabs = st.tabs(["Upload PMSG Approval of feasibility", "Generated Agreement"])
        with tabs[0]:
            render_upload_feasibility_tab()
        with tabs[1]:
            render_generated_agreements_tab()


# ---------------------------
# UI helpers
# ---------------------------

def render_dashboard():
    st.subheader("Dashboard")

    # Light styling for nicer cards/sections
    st.markdown(
        """
        <style>
        .dash-section h3, .dash-section h4 { margin-top: 0.5rem; }
        .dash-card { padding: 0.6rem 0.8rem; border: 1px solid #e5e7eb; border-radius: 10px; background: #ffffff; }
        .dash-muted { color:#6b7280; }
        </style>
        """,
        unsafe_allow_html=True,
    )

    # Load data for dashboard with broader columns
    conn = get_conn()
    try:
        df = pd.read_sql_query(
            """
            SELECT id, customer_name, product, date_of_quotation, quotation_no, pdf_path, created_at
            FROM invoices
            ORDER BY id DESC
            """,
            conn,
        )
    finally:
        conn.close()

    if df.empty:
        st.info("No invoices yet.")
        return

    # Normalize dates
    def to_date(x):
        try:
            return datetime.strptime(str(x), "%Y-%m-%d").date()
        except Exception:
            try:
                return datetime.fromisoformat(str(x)).date()
            except Exception:
                return None
    df["date_of_quotation"] = df["date_of_quotation"].apply(to_date)

    today = datetime.now().date()
    yesterday = today - timedelta(days=1)
    week_start = today - timedelta(days=today.weekday())  # Monday
    week_end = week_start + timedelta(days=6)
    month_start = today.replace(day=1)
    # compute month_end
    if month_start.month == 12:
        next_month_start = month_start.replace(year=month_start.year + 1, month=1)
    else:
        next_month_start = month_start.replace(month=month_start.month + 1)
    month_end = next_month_start - timedelta(days=1)

    # Controls
    c1, c2, c3 = st.columns([2, 2, 2])
    with c1:
        range_opt = st.radio(
            "Date Range",
            options=["Today", "Yesterday", "This Week", "This Month", "Custom"],
            horizontal=True,
        )
    with c2:
        sort_opt = st.selectbox("Sort by Date", options=["Newest first", "Oldest first"], index=0)
    with c3:
        # Placeholder for custom date control below
        pass

    start_date, end_date = None, None
    if range_opt == "Today":
        start_date, end_date = today, today
    elif range_opt == "Yesterday":
        start_date, end_date = yesterday, yesterday
    elif range_opt == "This Week":
        start_date, end_date = week_start, week_end
    elif range_opt == "This Month":
        start_date, end_date = month_start, month_end
    else:
        # Custom: let user pick a date or date range
        dr = st.date_input("Select date or range", value=(week_start, week_end))
        if isinstance(dr, tuple) and len(dr) == 2:
            start_date, end_date = dr
        else:
            start_date = dr
            end_date = dr

    # Filter
    df_valid = df.dropna(subset=["date_of_quotation"]).copy()
    if start_date and end_date:
        mask = (df_valid["date_of_quotation"] >= start_date) & (df_valid["date_of_quotation"] <= end_date)
        df_view = df_valid[mask]
    else:
        df_view = df_valid

    # Metrics row
    total_invoices = len(df)
    # Count rows that have a non-empty pdf_path (fix parentheses to sum the boolean mask)
    total_with_pdf = int(((df["pdf_path"].notna()) & (df["pdf_path"] != "")).sum()) if "pdf_path" in df.columns else 0
    selected_count = len(df_view)
    unique_customers = df["customer_name"].nunique()

    m1, m2, m3, m4 = st.columns(4)
    with m1:
        st.metric("Total Invoices", total_invoices)
    with m2:
        st.metric("Invoices in Range", selected_count)
    with m3:
        st.metric("Unique Customers", int(unique_customers))
    with m4:
        st.metric("With PDF", total_with_pdf)

    # Agreements & Feasibility section (date-wise, same range controls)
    st.markdown("### Agreements & Feasibility")

    # Helper: robust date parse for agreement.created_at (iso) and agreement.date (varied)
    def _to_date_any(x):
        if x is None:
            return None
        t = str(x).strip()
        for fmt in ("%Y-%m-%d", "%d-%m-%Y", "%Y/%m/%d", "%d/%m/%Y", "%Y-%m-%dT%H:%M:%S"):
            try:
                return datetime.strptime(t, fmt).date()
            except Exception:
                pass
        try:
            return datetime.fromisoformat(t).date()
        except Exception:
            return None

    # Agreements created count by created_at within range
    conn = get_conn()
    try:
        df_ag = pd.read_sql_query(
            """
            SELECT id, created_at FROM agreements ORDER BY id DESC
            """,
            conn,
        )
    finally:
        conn.close()
    if not df_ag.empty:
        df_ag["created_date"] = df_ag["created_at"].apply(_to_date_any)
        df_ag_valid = df_ag.dropna(subset=["created_date"]).copy()
        if start_date and end_date:
            mask_ag = (df_ag_valid["created_date"] >= start_date) & (df_ag_valid["created_date"] <= end_date)
            df_ag_view = df_ag_valid[mask_ag]
        else:
            df_ag_view = df_ag_valid
        agreements_in_range = len(df_ag_view)
    else:
        agreements_in_range = 0

    # Feasibility uploads count from filesystem (by file modified time)
    feas_in_range = 0
    try:
        if os.path.isdir(FEASIBILITY_DIR):
            for name in os.listdir(FEASIBILITY_DIR):
                if not name.lower().endswith(".pdf"):
                    continue
                fpath = os.path.join(FEASIBILITY_DIR, name)
                try:
                    dt = datetime.fromtimestamp(os.path.getmtime(fpath)).date()
                    if start_date and end_date:
                        if start_date <= dt <= end_date:
                            feas_in_range += 1
                    else:
                        feas_in_range += 1
                except Exception:
                    continue
    except Exception:
        pass

    cfa, cag = st.columns(2)
    with cfa:
        st.metric("Feasibility Uploaded (range)", feas_in_range)
    with cag:
        st.metric("Agreements Created (range)", agreements_in_range)

    # Quick glance mini-counters
    def count_in(d0, d1):
        return int(((df_valid["date_of_quotation"] >= d0) & (df_valid["date_of_quotation"] <= d1)).sum())
    q1, q2, q3, q4 = st.columns(4)
    with q1:
        st.caption("Today")
        st.write(count_in(today, today))
    with q2:
        st.caption("Yesterday")
        st.write(count_in(yesterday, yesterday))
    with q3:
        st.caption("This Week")
        st.write(count_in(week_start, week_end))
    with q4:
        st.caption("This Month")
        st.write(count_in(month_start, month_end))

    # Charts
    st.markdown("### Charts")
    ch1, ch2 = st.columns(2)
    # Invoices by Product (bar)
    with ch1:
        st.markdown("#### Invoices by Product")
        try:
            prod_counts = (
                df_view.groupby("product").size().reset_index(name="count").sort_values("count", ascending=False)
            )
            if not prod_counts.empty:
                prod_counts = prod_counts.set_index("product")
                st.bar_chart(prod_counts["count"])
            else:
                st.caption("No data for selected range.")
        except Exception:
            st.caption("Unable to render chart.")
    # Invoices over time (line)
    with ch2:
        st.markdown("#### Invoices over Time")
        try:
            daily_counts = (
                df_view.groupby("date_of_quotation").size().reset_index(name="count").sort_values("date_of_quotation")
            )
            if not daily_counts.empty:
                daily_counts = daily_counts.set_index("date_of_quotation")
                st.line_chart(daily_counts["count"])
            else:
                st.caption("No data for selected range.")
        except Exception:
            st.caption("Unable to render chart.")

    # Removed Top Customers section as requested

    # Additional: Distribution by Day of Week
    st.markdown("#### Distribution by Day of Week")
    try:
        dts = pd.to_datetime(df_view["date_of_quotation"], errors="coerce")
        dow = dts.dt.day_name()
        order = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
        dow_counts = (
            dow.value_counts().reindex(order).fillna(0).astype(int)
        )
        if dow_counts.sum() > 0:
            st.bar_chart(dow_counts)
        else:
            st.caption("No data for selected range.")
    except Exception:
        st.caption("Unable to render chart.")

    # Invoices section intentionally removed

def render_create_form(
    prefill: Optional[Dict[str, str]] = None,
    edit_id: Optional[int] = None,
    allowed_products: Optional[List[str]] = None,
    form_title: Optional[str] = None,
    key_ns: Optional[str] = None,
):
    conn = get_conn()
    try:
        qno_preview = next_quotation_no(conn) if edit_id is None else None
    finally:
        conn.close()

    # establish a namespace for widget keys to avoid collisions across tabs/forms
    ns = key_ns or (f"edit_{edit_id}" if edit_id is not None else "new")

    if edit_id is None:
        st.subheader(form_title or "Create New Invoice")
    else:
        st.subheader("Edit Invoice")

    # Date of Quotation & Validity
    # Create mode: allow picking date; Edit mode: lock and use stored values
    if edit_id is None:
        date_default = datetime.now().date()
        if prefill and prefill.get("date_of_quotation"):
            try:
                date_default = datetime.strptime(prefill["date_of_quotation"], "%Y-%m-%d").date()
            except Exception:
                pass
        date_of_quotation = st.date_input(
            "Date of Quotation",
            value=date_default,
            key=f"{ns}_doq",
        )
        validity_date = date_of_quotation + timedelta(days=30)
    else:
        # Parse existing values to date objects for consistent downstream formatting
        doq_str = (prefill.get("date_of_quotation") if prefill else None) or datetime.now().date().isoformat()
        try:
            date_of_quotation = datetime.strptime(str(doq_str), "%Y-%m-%d").date()
        except Exception:
            date_of_quotation = datetime.now().date()
        val_str = (prefill.get("validity_date") if prefill else None)
        if val_str:
            try:
                validity_date = datetime.strptime(str(val_str), "%Y-%m-%d").date()
            except Exception:
                validity_date = date_of_quotation + timedelta(days=30)
        else:
            validity_date = date_of_quotation + timedelta(days=30)

    with st.form(key=f"{ns}_form"):
        product_options = allowed_products if allowed_products else PRODUCT_OPTIONS
        default_index = 0
        if prefill and prefill.get("product") in product_options:
            default_index = product_options.index(prefill.get("product"))
        product = st.selectbox("Product & Service", options=product_options, index=default_index, key=f"{ns}_product")
        customer_name = st.text_input("Customer Name", value=(prefill.get("customer_name") if prefill else ""), key=f"{ns}_customer")
        mobile = st.text_input("Mobile Number", value=(prefill.get("mobile") if prefill else ""), key=f"{ns}_mobile")
        location = st.text_input("Location", value=(prefill.get("location") if prefill else ""), key=f"{ns}_location")
        city = st.text_input("City", value=(prefill.get("city") if prefill else ""), key=f"{ns}_city")
        state = st.text_input("State", value=(prefill.get("state") if prefill else ""), key=f"{ns}_state")
        pincode = st.text_input("Pincode", value=(prefill.get("pincode") if prefill else ""), key=f"{ns}_pincode")
        staff_name = st.text_input("Staff Name (kept only in DB)", value=(prefill.get("staff_name") if prefill else ""), key=f"{ns}_staff")
        # Show computed values inside the form (read-only)
        st.text_input("Date of Quotation", value=date_of_quotation.isoformat(), disabled=True, key=f"{ns}_doq_ro")
        st.text_input("Validity (auto)", value=validity_date.isoformat(), disabled=True, key=f"{ns}_validity_ro")

        if edit_id is None and qno_preview:
            st.text_input("Quotation No (auto)", value=qno_preview, disabled=True, key=f"{ns}_qno_preview")
        elif edit_id is not None and prefill and prefill.get("quotation_no"):
            st.text_input("Quotation No", value=prefill.get("quotation_no"), disabled=True, key=f"{ns}_qno_edit")

        submit_label = "Update Invoice" if edit_id is not None else "Create Invoice"
        submitted = st.form_submit_button(submit_label)

    if submitted:
        data = {
            "product": product,
            "customer_name": customer_name.strip(),
            "mobile": str(mobile).strip(),
            "location": location.strip(),
            "city": city.strip(),
            "state": state.strip(),
            "pincode": str(pincode).strip(),
            "staff_name": staff_name.strip(),
            "date_of_quotation": date_of_quotation.isoformat(),
            "validity_date": validity_date.isoformat(),
        }
        # Minimal progress UI (non-intrusive)
        prog = st.progress(0, text="Starting…")
        status = st.empty()
        try:
            prog.progress(10, text="Processing invoice…")
            status.write("Generating files…")
            # choose template based on selected product
            template_path = _template_for_product(product)
            if edit_id is None:
                docx_path, pdf_path = create_invoice(data, template_path)
                prog.progress(70, text="Finalizing creation…")
                st.success("Invoice created successfully.")
                st.toast(f"Saved invoice {qno_preview}", icon="✅")
            else:
                docx_path, pdf_path = edit_invoice(edit_id, data, template_path)
                prog.progress(70, text="Finalizing update…")
                st.success("Invoice updated successfully.")
                st.toast("Invoice updated", icon="✏️")
            if pdf_path and os.path.exists(pdf_path):
                # Inline preview of the generated PDF
                _render_pdf_preview(pdf_path, height=480)
                # Actions: Download + Share side-by-side
                cdl, csh = st.columns([1, 1])
                with cdl:
                    with open(pdf_path, "rb") as f:
                        prog.progress(90, text="Preparing download…")
                        st.download_button(
                            "⬇️  Download",
                            data=f.read(),
                            file_name=os.path.basename(pdf_path),
                            mime="application/pdf",
                            use_container_width=True,
                        )
                with csh:
                    _render_mobile_share_button(pdf_path, os.path.basename(pdf_path))
                prog.progress(100, text="Done")
            else:
                prog.progress(100, text="Completed (PDF unavailable)")
                if docx_path and isinstance(docx_path, str) and os.path.exists(docx_path):
                    st.warning("PDF conversion failed. Download the DOCX and export to PDF using Word/LibreOffice. You can also install LibreOffice or MS Word to enable automatic PDF generation.")
                    with open(docx_path, "rb") as f:
                        st.download_button(
                            "⬇️  Download DOCX",
                            data=f.read(),
                            file_name=os.path.basename(docx_path),
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                        )
                    st.caption("Tip: On Windows, installing MS Word usually enables automatic PDF conversion via docx2pdf. Alternatively, install LibreOffice and set environment variable LIBREOFFICE_PATH to the soffice.exe.")
                else:
                    st.warning("PDF conversion failed or Word is not available. Please try again on a system with MS Word or LibreOffice installed.")
        except Exception as e:
            prog.progress(100, text="Failed")
            st.error(f"Failed to process invoice: {e}")


def render_search_tab():
    st.subheader("Search Invoice")
    # Process delete action if triggered via query param
    _handle_delete_via_query()

    df = load_invoices()
    if df.empty:
        st.info("No invoices yet.")
        return

    # Filters
    c1, c2, c3 = st.columns([2, 2, 2])
    with c1:
        f_name = st.text_input("Filter by Customer Name")
    with c2:
        f_mobile = st.text_input("Filter by Mobile Number")
    with c3:
        f_product = st.selectbox("Filter by Product", options=["All"] + PRODUCT_OPTIONS, index=0)

    mask = pd.Series([True] * len(df))
    if f_name:
        mask &= df["customer_name"].str.contains(f_name, case=False, na=False)
    if f_mobile:
        mask &= df["mobile"].astype(str).str.contains(f_mobile, case=False, na=False)
    if f_product != "All":
        mask &= df["product"] == f_product

    filtered = df[mask].reset_index(drop=True)

    # Handle action links via query params (modern API)
    qp = st.query_params
    def _first(v):
        if v is None:
            return None
        if isinstance(v, list):
            return v[0] if v else None
        return v
    action = _first(qp.get("action"))
    action_id = _first(qp.get("id"))
    if action and action_id:
        try:
            rid = int(action_id)
            if action == "preview":
                st.session_state["preview_id"] = rid
            elif action == "edit":
                st.session_state["selected_edit_id"] = rid
                st.session_state.pop("preview_id", None)
            elif action == "delete":
                delete_invoice(rid)
                st.success("Deleted.")
            # Clear params to avoid repeat on next runs (no extra rerun here)
            st.query_params.clear()
        except Exception:
            pass

    # Toggle for mobile card view
    mobile_view = st.toggle("Mobile card view", value=True, key="mobile_card_toggle")

    # Light CSS for compact icon buttons and spacing
    st.markdown(
        """
        <style>
        .card-header {display:flex; justify-content:space-between; align-items:center;}
        .card-title {font-weight:600; margin: 0;}
        .meta {color:#6b7280; font-size:12px; margin: 0;}
        /* Make Streamlit buttons look compact */
        .stButton>button {padding: 0.35rem 0.6rem; border-radius:999px; font-size:13px;}
        .stDownloadButton>button {padding: 0.35rem 0.6rem; border-radius:999px; font-size:13px;}
        /* Force inline horizontal layout for buttons even on mobile */
        .stButton, .stDownloadButton {display:inline-block !important; margin: 0 8px 8px 0 !important;}
        .stButton>button, .stDownloadButton>button {min-width: 36px; height: 36px;}
        .card-block {padding-top: 0.25rem;}
        .action-links {display:flex; align-items:center; gap: 10px; flex-wrap: nowrap; margin-bottom: 16px;}
        .action-links a {text-decoration:none; color:#374151; background:#f3f4f6; padding:6px 10px; border-radius:999px; font-size:13px; display:inline-flex; align-items:center; gap:6px;}
        .action-links a:hover {background:#e5e7eb;}
        </style>
        """,
        unsafe_allow_html=True,
    )

    # Smooth UX: preserve scroll position across reruns so actions feel inline without jump
    components.html(
        """
        <script>
        (function(){
          const KEY = 'search_invoice_scrollY';
          const y = sessionStorage.getItem(KEY);
          if (y) { try { window.scrollTo(0, parseInt(y)); } catch (e) {} }
          window.addEventListener('beforeunload', function(){
            try { sessionStorage.setItem(KEY, String(window.scrollY)); } catch(e) {}
          });
        })();
        </script>
        """,
        height=1,
    )

    if mobile_view:
        # Card layout per row (good on mobile)
        for _, row in filtered.iterrows():
            with st.container(border=True):
                pdf_path = row.get("pdf_path") if "pdf_path" in row else None

                # Header (no actions here to avoid vertical stacking on small screens)
                top_l, _ = st.columns([7, 3])
                with top_l:
                    st.markdown(
                        f"<p class='card-title'>{row['customer_name']}</p>",
                        unsafe_allow_html=True,
                    )
                    st.markdown(
                        f"<p class='meta'>Quotation No: {row['quotation_no']}</p>",
                        unsafe_allow_html=True,
                    )
                # Actions: force single horizontal row using 4 columns
                rid = int(row["id"])
                a1, a2, a3, a4, a5 = st.columns([1, 1, 1, 1, 1])
                with a1:
                    if st.button("👁️  Preview", key=f"m_prev_{rid}", use_container_width=True):
                        st.session_state["preview_id"] = rid
                        st.session_state.pop("selected_edit_id", None)
                with a2:
                    if pdf_path and os.path.exists(pdf_path):
                        try:
                            with open(pdf_path, "rb") as f:
                                st.download_button(
                                    "⬇️  Download",
                                    data=f.read(),
                                    file_name=os.path.basename(pdf_path),
                                    mime="application/pdf",
                                    key=f"m_dl_{rid}",
                                    use_container_width=True,
                                )
                        except Exception:
                            st.button("⬇️  Download", disabled=True, key=f"m_dl_dis_{rid}", use_container_width=True)
                    else:
                        st.button("⬇️  Download", disabled=True, key=f"m_dl_na_{rid}", use_container_width=True)
                with a3:
                    if st.button("✏️  Edit", key=f"m_edit_{rid}", use_container_width=True):
                        st.session_state["selected_edit_id"] = rid
                        st.session_state.pop("preview_id", None)
                with a4:
                    if pdf_path and os.path.exists(pdf_path):
                        _render_mobile_share_button(pdf_path, os.path.basename(pdf_path))
                    else:
                        st.button("Share PDF", disabled=True, key=f"m_share_na_{rid}", use_container_width=True)
                with a5:
                    if st.button("🗑️  Delete", key=f"m_del_{rid}", use_container_width=True):
                        delete_invoice(rid)
                        st.success("Deleted.")
                        st.rerun()

                # Compact details with View more (show all key fields)
                with st.expander("View details", expanded=False):
                    rec = fetch_full_record(rid) or {}
                    def v(key):
                        return rec.get(key, "")
                    st.markdown(f"**Quotation No (auto)**: {v('quotation_no')}")
                    st.markdown(f"**Product & Service**: {v('product')}")
                    st.markdown(f"**Customer Name**: {v('customer_name')}")
                    st.markdown(f"**Mobile Number**: {v('mobile')}")
                    st.markdown(f"**Location**: {v('location')}")
                    st.markdown(f"**City**: {v('city')}")
                    st.markdown(f"**State**: {v('state')}")
                    st.markdown(f"**Pincode**: {v('pincode')}")
                    st.markdown(f"**Staff Name (kept only in DB)**: {v('staff_name')}")
                    st.markdown(f"**Date of Quotation**: {v('date_of_quotation')}")
                    st.markdown(f"**Validity**: {v('validity_date')}")

                # Inline PDF preview (if chosen)
                if st.session_state.get("preview_id") == int(row["id"]) and pdf_path and os.path.exists(pdf_path):
                    _render_pdf_preview(pdf_path, height=480)
                    if st.button("Close preview", key=f"m_close_prev_{row['id']}"):
                        st.session_state.pop("preview_id", None)
                        st.rerun()

                # Inline share prompt removed; Share PDF button is now directly in the action row

                # Inline Edit panel (only for the selected card)
                if st.session_state.get("selected_edit_id") == int(row["id"]):
                    st.markdown("#### Edit Invoice")
                    prefill = fetch_full_record(int(row["id"])) or {}
                    c_cancel, _ = st.columns([1, 6])
                    with c_cancel:
                        if st.button("Close", key=f"close_edit_inline_{row['id']}"):
                            st.session_state.pop("selected_edit_id", None)
                            st.rerun()
                    render_create_form(prefill=prefill, edit_id=int(row["id"]))

        return

    # Desktop-like table layout with single Action column (previous behavior)
    # Header
    h1, h2, h3, h4, h5 = st.columns([2.5, 2.5, 2, 2, 2])
    with h1:
        st.markdown("**Customer**")
    with h2:
        st.markdown("**Product**")
    with h3:
        st.markdown("**Date of Quotation**")
    with h4:
        st.markdown("**Quotation No**")
    with h5:
        st.markdown("**Action**")

    # Rows
    for _, row in filtered.iterrows():
        rid = int(row["id"])
        pdf_path = row.get("pdf_path") if "pdf_path" in row else None
        c1, c2, c3, c4, c5 = st.columns([2.5, 2.5, 2, 2, 2])
        with c1:
            st.write(row["customer_name"])  
        with c2:
            st.write(row["product"]) 
        with c3:
            st.write(row["date_of_quotation"]) 
        with c4:
            st.write(row["quotation_no"]) 
        with c5:
            a1, a2, a3, a4, a5 = st.columns([1, 1, 1, 1, 1])
            with a1:
                if st.button("👁️", key=f"d_prev_{rid}", use_container_width=True):
                    st.session_state["preview_id"] = rid
                    st.session_state.pop("selected_edit_id", None)
            with a2:
                if pdf_path and os.path.exists(pdf_path):
                    with open(pdf_path, "rb") as f:
                        st.download_button("⬇️", f.read(), file_name=os.path.basename(pdf_path), key=f"d_dl_{rid}", use_container_width=True)
                else:
                    st.button("⬇️", disabled=True, key=f"d_dl_na_{rid}", use_container_width=True)
            with a3:
                if st.button("✏️", key=f"d_edit_{rid}", use_container_width=True):
                    st.session_state["selected_edit_id"] = rid
                    st.session_state.pop("preview_id", None)
                    st.rerun()
            with a4:
                if pdf_path and os.path.exists(pdf_path):
                    _render_mobile_share_button(pdf_path, os.path.basename(pdf_path))
                else:
                    st.button("Share PDF", disabled=True, key=f"d_share_na_{rid}", use_container_width=True)
            with a5:
                if st.button("🗑️", key=f"d_del_{rid}", use_container_width=True):
                    delete_invoice(rid)
                    st.success("Deleted.")
                    st.rerun()

        # Inline preview right under the targeted row (desktop view), same as mobile behavior
        if st.session_state.get("preview_id") == rid and pdf_path and os.path.exists(pdf_path):
            with st.container():
                _render_pdf_preview(pdf_path, height=480)
                if st.button("Close preview", key=f"d_close_preview_{rid}"):
                    st.session_state.pop("preview_id", None)
                    st.rerun()

        # Inline share prompt removed; Share PDF is now directly in the action row (desktop view)

        # Inline edit panel right under the targeted row (desktop view), same as mobile behavior
        if st.session_state.get("selected_edit_id") == rid:
            with st.container():
                st.markdown("#### Edit Invoice")
                prefill = fetch_full_record(rid) or {}
                c_cancel, _ = st.columns([1, 6])
                with c_cancel:
                    if st.button("Close", key=f"d_close_edit_{rid}"):
                        st.session_state.pop("selected_edit_id", None)
                        st.rerun()
                render_create_form(prefill=prefill, edit_id=rid)

    # Global preview/edit panels are intentionally removed for desktop table view to keep UI inline per row


def fetch_full_record(inv_id: int) -> Optional[Dict[str, str]]:
    conn = get_conn()
    try:
        cur = conn.cursor()
        cur.execute(
            "SELECT id, quotation_no, product, customer_name, mobile, location, city, state, pincode, staff_name, date_of_quotation, validity_date FROM invoices WHERE id=?",
            (inv_id,),
        )
        r = cur.fetchone()
        if not r:
            return None
        return {
            "id": r[0],
            "quotation_no": r[1],
            "product": r[2],
            "customer_name": r[3],
            "mobile": r[4],
            "location": r[5],
            "city": r[6],
            "state": r[7],
            "pincode": r[8],
            "staff_name": r[9],
            "date_of_quotation": r[10],
            "validity_date": r[11],
        }
    finally:
        conn.close()


if __name__ == "__main__":
    main()
